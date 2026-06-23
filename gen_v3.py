"""
v3 讲解稿 — 技术详解版。中文文本先定义后写入，避免引号冲突。
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

def SF(run, cn, en, size, bold=False, color=None):
    run.font.size = Pt(size); run.bold = bold; run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rF = rPr.find(qn('w:rFonts'))
    if rF is None: rF = OxmlElement('w:rFonts'); rPr.insert(0, rF)
    rF.set(qn('w:eastAsia'), cn); rF.set(qn('w:ascii'), en); rF.set(qn('w:hAnsi'), en)
    if color: run.font.color.rgb = color

def T(text, level=1):
    h = doc.add_heading(text, level=level); h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in h.runs: SF(r, '黑体', 'Arial', {0:22,1:16,2:14}.get(level,14), True)

def P(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); SF(r, '宋体', 'Consolas', 12)

def C(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74); p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); SF(r, '宋体', 'Consolas', 10.5, color=RGBColor(0x16,0x77,0xFF))

# ========== 文本内容定义（存为变量，避免行内引号冲突） ==========

t_cover_title = "NETMAXXX 学生管理系统"
t_cover_sub   = "微信小程序演示讲解稿（技术详解版 v3）"

# 一、项目概述
t1a = "大家好，我今天展示的项目是 NETMAXXX 学生管理系统。这是一个基于微信小程序平台开发的、面向高校师生日常管理场景的综合信息管理系统。系统支持学生和教师两种角色，涵盖学生档案管理、课堂签到、请假审批、数据统计分析和 AI 智能对话等完整的校园管理功能模块。"

t1b = "在技术选型上，前端采用微信原生框架进行开发，没有使用 uni-app 或 Taro 等第三方跨端框架，因此可以最大程度地调用微信小程序的底层硬件能力——包括 GPS 定位、摄像头扫码、地图渲染、文件系统操作等。每个页面由四种文件组成：WXML 模板文件负责界面结构，使用微信自己的组件体系如 view、text、image、input、picker、map 等；WXSS 样式文件负责视觉效果，语法兼容 CSS，支持 flexbox 弹性布局、CSS 变量、关键帧动画等现代特性；JS 逻辑文件使用 Page() 或 Component() 构造器注册，通过 setData 方法驱动视图更新；JSON 配置文件负责页面级设置，如导航栏标题、下拉刷新开关、引用的自定义组件等。"

t1c = "前端样式体系的设计是本次开发的一个重点。在全局样式文件 app.wxss 中，我们使用 CSS 自定义属性（也叫 CSS 变量）定义了一套完整的设计令牌——包括 --color-primary 品牌蓝色 #1677FF、--color-success 成功绿色 #00B42A、--color-warning 警告橙色 #FF7D00、--color-danger 危险红色 #F53F3F 等语义化颜色变量；--radius-sm/md/lg 三级圆角变量（4rpx/8rpx/12rpx）；--spacing-xs/sm/md/lg 四级间距变量（8rpx~32rpx）；--font-xs/sm/md/lg/xl 五级字号变量（22rpx~36rpx）。这些变量统一声明在 page 选择器上，所有子页面和组件通过 var() 函数引用，实现了全局主题的一致性。当需要调整主题色时，只需修改 app.wxss 中的一处定义即可全局生效。"

t1d = "在此基础上，app.wxss 还定义了一套原子化的工具类，包括 flex 布局工具类（.flex/.flex-center/.flex-between/.flex-1）、文本工具类（.text-center/.text-muted/.text-primary/.text-danger）、间距工具类（.mt-sm/.mt-md/.p-sm/.p-md 等）、以及 .tag 标签体系。.tag 标签通过五个颜色变体类名（tag-primary/tag-success/tag-warning/tag-danger/tag-info）覆盖了系统中所有标签场景——性别标签、请假类型标签、签到状态标签、审批状态标签全部复用同一套 CSS 类，保证了视觉一致性。按钮系统同样通过 .btn 基类加 .btn-primary/.btn-success/.btn-danger/.btn-outline/.btn-sm/.btn-block 等修饰类组合使用，实现了灵活的按钮样式组合。"

t1e = "后端基于 Django 5.2 框架构建，使用 MySQL 数据库，通过 natapp 内网穿透将本地服务映射到公网。前后端通过 RESTful API 通信，JSON 格式传参，JWT 双 Token 机制做身份认证。下面我将按照功能模块逐一介绍技术实现和样式设计。"

# 二、登录与注册
t2a = "登录页是整个小程序的入口，也是用户对产品视觉风格的第一印象。页面采用全屏渐变背景——使用 CSS 的 linear-gradient 函数，从顶部的 #1677FF 品牌蓝渐变到底部的 #69b1ff 浅蓝，方向为 135 度斜向过渡，配合 min-height: 100vh 确保覆盖整个屏幕高度。内容是垂直居中的 flex 布局，由 Logo 标题区、模式切换选项卡和表单卡片三部分组成。"

t2b = "模式切换选项卡（登录/注册）使用了半透明磨砂玻璃风格。外层容器设置 background: rgba(255,255,255,0.2) 半透明白色背景，border-radius: 40rpx 大圆角，padding: 6rpx 内边距形成胶囊形状。内部的每个 tab 项通过 .mode-tab-active 类名标识选中态——选中时 background 变为纯白色、color 变为品牌蓝、font-weight 加粗、box-shadow 添加轻微阴影实现浮起效果。切换过程使用 transition: all 0.3s 实现平滑过渡动画。角色选择按钮（学生入口/教师入口）采用类似的设计。"

t2c = "表单卡片是一个白色圆角容器，设置 background: #fff、border-radius: 24rpx 大圆角和 box-shadow: 0 8rpx 32rpx rgba(0,0,0,0.1) 深阴影，从渐变背景中浮起形成层次感。输入框在卡片内使用浅灰背景 #f7f8fa 以区分于卡片白色底，聚焦时通过 :focus 伪类切换为白色背景加蓝色边框，给予用户明确的交互反馈。"

t2d = "顶部消息提示条使用 position: fixed 固定定位在导航栏下方，通过 @keyframes slideDown 动画从上方滑入——from 状态的 transform: translateY(-100%) 到 to 状态的 translateY(0)，配合 0.3s ease 缓动函数。错误消息使用浅红底 #fff0f0 配红色文字和红色边框，成功消息使用浅绿底 #e8ffea 配绿色文字。"

t2e = "在密码输入框方面，我们增加了一个显示/隐藏密码的切换控件。它由一个自定义复选框和一个文字标签组成，复选框是一个 36rpx 见方的方块，未选中时灰色边框白色背景，选中后通过 .checked 类名切换为蓝色背景加白色对勾。点击时通过 bindtap 事件翻转 showPassword 布尔值，input 组件的 password 属性绑定 {{!showPassword}} 表达式——当 showPassword 为 true 时 password 属性值为 false，密码即以明文显示。"

# 三、角色权限与TabBar
t3a = "系统的角色权限通过微信小程序的自定义 TabBar 机制实现。与原生 TabBar 的固定配置不同，自定义 TabBar 是一个独立的自定义组件，定义在 custom-tab-bar 目录下，包含 index.js、index.wxml、index.wxss 和 index.json 四个文件。在 app.json 中设置 custom: true 开启自定义模式后，微信客户端不再渲染原生底部导航，而是由开发者在每个 Tab 页面的 onShow 生命周期中手动调用 this.getTabBar() 获取组件实例来驱动更新。"

t3b = "TabBar 组件的核心逻辑在 updateTabs 方法中。该方法读取 app.globalData.userRole 获取当前用户角色，然后从一个预定义的 allTabs 数组中过滤出当前角色可见的标签项。每个标签项定义四个属性：pagePath 页面路由、text 显示文字、iconPath 图标路径、role 可见角色（teacher/student/all）。教师端显示学生信息、数据中心、签到、我的四个标签，学生端显示签到、请假、我的三个标签。标签切换通过 wx.switchTab API 实现。"

t3c = "在样式上，TabBar 采用底部固定定位，通过 wx:if 控制显隐——当签到地图或 AI 对话全屏展开时，页面通过 getTabBar().setData({ tabBarHidden: true }) 隐藏底部导航，避免遮挡内容区域。每个标签项使用 flex 纵向排列，图标在上文字在下，选中态的图标和文字通过 .icon-active 和 .text-active 类名切换为品牌蓝色。"

t3d = "权限控制的核心在后端。以学生列表接口为例——views.py 中的 get_students 函数会先从请求头中提取 Bearer Token 并解析出用户角色。如果角色是 student，则只返回 user_id 匹配的单条学生记录；如果是 teacher，则返回全量数据。这种在视图函数层面基于角色的数据过滤，从根本上保证了数据安全，前端只做 UI 层面的菜单显隐，真正的权限防线在后端。"

# 四、学生信息管理
t4a = "学生信息管理是系统最基础的数据模块。数据库层面，Student 表包含八个字段，学号 sno 使用 IntegerField 并设置 primary_key=True 作为主键，Gender 字段通过 choices 参数限制可选值为男或女。Django Model 的 Meta 内部类中设置 db_table = 'Student' 指定表名，managed = True 表示由 Django 管理迁移。前端表单通过 validator.js 中的正则表达式进行多层校验：学号必须匹配 95 开头的五位数格式，姓名必须匹配 2 到 5 个汉字，手机号匹配 11 位标准手机号格式。"

t4b = "学生列表页采用 flex 水平布局的搜索栏——input 组件设置 flex: 1 占据剩余空间，配合查询和全部两个按钮。统计信息栏使用 .text-muted 灰色文字显示总记录数和当前页码。每张学生卡片是一个水平 flex 容器，左侧是圆形头像区域——使用 image 组件配合 mode='aspectFill' 保持比例裁剪填充，无头像时显示取姓名首字的文字占位符，占位符使用蓝色背景白色文字的圆形设计。中间信息区展示姓名、性别标签和学号，性别标签通过条件类名动态绑定实现蓝男红女的颜色区分。右侧操作按钮区仅在教师身份时通过 wx:if 渲染。卡片左侧复选框选中时通过 .checkbox-checked 类名显示蓝色背景加白色对勾。"

t4c = "Excel 导入功能通过微信的 wx.chooseMessageFile API 实现文件选择——设置 count: 1 限制单文件，type: 'file' 指定选择文件，extension: ['xlsx','xls'] 限制文件类型。文件选择后通过封装的 api.uploadExcel 方法上传——底层使用 wx.uploadFile API 以 multipart/form-data 格式传输。后端使用 openpyxl 库解析，通过 read_excel_dict 函数智能识别表头格式——先读取第一行判断是否包含学号、姓名等关键字，有表头则按列名映射，没有则按固定列序解析，兼容了多种 Excel 模板格式。"

# 五、签到功能
t5_intro = "签到功能是本系统技术实现最丰富的模块，深度融合了微信小程序的 GPS 定位和摄像头扫码两大硬件能力。整个模块在单一页面 checkin.js 中通过最外层 wx:if 条件渲染实现了教师和学生两套完全不同的 UI。接下来我按照教师发布任务、学生扫码签到、学生位置签到、教师考勤管理四个环节逐一介绍技术实现和样式设计。"

t5a = "教师的签到发布表单使用了三种小程序原生表单组件。签到标题使用 input 组件，通过 bindinput 事件将输入内容实时同步到 taskTitle 数据字段。签到方式选择使用 picker 组件——mode 设为 selector 单列选择模式，range 绑定 checkinTypes 数组（包含二维码签到和位置签到两个字符串），value 绑定当前选中索引 checkinTypeIndex，用户选择后触发 bindchange 事件回调。位置签到专属的范围滑块使用 slider 组件——min 设为 10 表示最小 10 米，max 设为 2000 表示最大 2000 米，step 设为 10 步进，activeColor 设为品牌蓝 #1677FF 使已滑过的轨道呈现蓝色，block-size 设为 20 控制滑块圆钮大小。滑块下方显示提示文字，X 动态绑定 rangeValue。"

t5b = "位置签到需要 GPS 定位。点击发布按钮前先调用 _getLocation 方法获取教师坐标。该方法核心是一个三层权限处理逻辑——使用 wx.getSetting 查询 scope.userLocation 的授权状态。关键设计在于三种状态的分支处理：状态为 true 已授权则直接调用 wx.getLocation；状态为 false 表示用户曾经明确拒绝，此时不直接调用定位（否则会静默失败），而是弹出 wx.showModal 对话框引导用户跳转到系统权限管理页；状态为 undefined 表示首次请求，直接调用定位，微信会自动弹出授权弹窗。定位调用使用 type: 'gcj02' 国测局坐标系——这是国内地图服务使用的加密坐标系，isHighAccuracy: true 开启 GPS 高精度模式，highAccuracyExpireTime: 15000 表示 15 秒超时后自动降级为网络定位以节省电量。"

t5c = "学生扫码签到通过微信的 wx.scanCode API 实现。设置 onlyFromCamera: true 表示仅允许从摄像头取景扫描，禁止从相册选取二维码图片。这个限制是有意为之的设计——它确保了物理场景属性：学生必须在课堂上实际对准教师展示的二维码才能签到，无法通过相册中的截图远程作弊。扫描成功后在 success 回调的 res.result 中获得解码后的签到的码字符串，将其与 task_id 一起提交到 checkin/sign/ 接口。"

t5d = "位置签到是技术实现最复杂的功能模块，核心依赖微信小程序的原生 map 组件。map 是微信提供的原生组件，底层由腾讯地图提供瓦片数据和渲染引擎，在小程序中有最高的渲染层级，因此可以实现流畅的地图交互。在 map 组件的属性配置方面：latitude 和 longitude 确定地图中心点经纬度——优先使用教师坐标为中心；scale 属性设为 17，对应大约 50 米的比例尺级别。markers 是标记点数组，学生端包含两个标记点：id 为 0 的红色标记表示我的位置，使用 32x32 像素较大尺寸，callout 属性设置气泡提示——content 显示位置名称、color 为红色 #E74C3C、display 设为 ALWAYS 使气泡始终显示、borderRadius 设为 8 圆角、padding 设为 6 内边距使气泡美观；label 属性设置文字标签。id 为 1 的蓝色标记表示教师位置，使用 28x28 像素略小尺寸。"

t5e = "circles 是地图上的圆形覆盖物数组，用于可视化签到有效范围。每个圆由 latitude/longitude 圆心坐标（设为教师位置）、radius 半径（单位米）、color 边框颜色（#1677FF66，最后两位是十六进制透明度）、fillColor 填充颜色（#1677FF1A，更高透明度）和 strokeWidth 边框宽度（2 像素）组成。通过半透明蓝色圆环，学生可以直观看到自己的标记点是否在有效签到范围内。"

t5f = "地图界面的布局采用全屏覆盖设计。外层容器使用 position: fixed 固定定位覆盖整个屏幕，z-index 设为较高值确保在所有内容之上。地图区域占据上方大部分空间。图例栏使用绝对定位在左上角，半透明白色背景加圆角和阴影，每个图例项是水平 flex 排列的小色块加文字说明——红色圆点标识教师位置、蓝色圆点标识学生位置，空心圆表示签到范围。底部操作面板使用白色背景加大顶部圆角营造从底部弹出的卡片效果，面板包含标题、提示文字和取消/确认两个按钮。"

t5g = "教师端查看学生签到位置时使用同样的 map 组件，但标记点数量更多。学生标记用较小尺寸（24x24 像素）与教师标记（32x32 像素）区分，callout 的 display 设为 BYCLICK（点击时显示），label 中显示数字序号。点击标记时触发 bindmarkertap 事件，从 e.detail.markerId 获取标记 ID，减 1 后从 locationList 数组中取出学生数据，用 wx.showToast 轻提示显示学生姓名和学号。"

t5h = "考勤详情面板是一个全屏对话框，使用 .dialog-mask 半透明黑色遮罩（background: rgba(0,0,0,0.5)）加 .dialog-box 白色圆角卡片实现。卡片标题区展示任务名称和三个统计数字——已签人数（绿色）、未签人数（红色）、总人数（灰色）。列表区域使用 scroll-view 设置 max-height: 60vh 限制最大高度并启用滚动。每行展示学生姓名学号、签到状态标签和操作按钮。未签到学生显示绿色补签按钮，已签到学生显示红色删除按钮，操作后自动刷新考勤数据。补签功能的设计解决了学生因手机故障等特殊情况无法正常签到的实际问题。"

# 六、请假功能
t6a = "请假管理实现了从申请提交、状态筛选到教师审批的完整业务闭环。Leave 表采用 12 个字段的冗余设计——student_id、student_name、student_no 三个字段同时存储学生标识信息，虽然牺牲了数据库范式，但使教师端查询时无需 JOIN Student 表即可直接获取学生姓名，显著提升了查询性能。status 字段使用字符串枚举 pending/approved/rejected 表示三态流转。审批操作同时记录 teacher_id、approval_comment 和 approval_time，形成完整的审批审计链。"

t6b = "学生端请假页面采用卡片折叠式设计。顶部申请表单卡片使用 .form-card 样式——白色背景、24rpx 大圆角、32rpx 内边距、轻微阴影。卡片头部是水平 flex 布局，左侧标题，右侧切换按钮文字通过三元表达式动态显示收起或展开。表单区域的显示隐藏使用 wx:if 而非 hidden 属性——wx:if 会完全销毁和重建 DOM，在表单收起时清空所有输入状态；hidden 仅仅是视觉隐藏，表单数据仍然保留。这个选择是有意为之的：用户收起表单意味着放弃当前输入，使用 wx:if 可以确保下次展开时获得全新的表单状态。"

t6c = "请假类型使用 picker mode='selector' 选择器。开始日期和结束日期使用 picker mode='date' 日期选择器——微信原生日期选择器提供了滚轮式的日期选择界面，支持年月日三级联动。关键设计在于结束日期选择器设置了 start 属性绑定开始日期的值，微信原生日期选择器会将早于 start 值的日期选项自动置灰不可选，从 UI 交互层面杜绝了结束日期早于开始日期的非法输入。"

t6d = "请假原因使用 textarea 多行文本组件——与单行 input 不同，textarea 支持换行和多行展示，适合输入较长的请假原因说明。组件设置 maxlength='500' 限制最大字符数，字符计数器 .char-count 使用绝对定位在右下角显示实时统计。textarea 的样式设置 min-height: 160rpx 保证足够的输入区域高度，边框和圆角与 input 保持一致形成统一的表单视觉。"

t6e = "请假记录列表采用卡片式布局，每条记录是一个 .leave-item 垂直 flex 容器。顶部水平展示两个 tag 标签——请假类型（病假橙色/事假灰色）和审批状态（待审批橙色/已批准绿色/已拒绝红色），标签系统全部复用全局定义的 .tag 样式体系。底部水平 flex 展示申请时间和操作链接——查看和删除，使用蓝色文字加下划线样式模拟超链接。"

t6f = "教师端请假管理页使用了一个独特的技术——wxs（WeiXin Script）视图层脚本。wxs 是微信小程序特有的脚本语言，运行在视图层（渲染层）而非逻辑层（JS 线程），可以直接在 WXML 模板中引入并调用其中的函数。页面引用的 wxs 文件导出了 isSelected 函数，用于判断某条请假记录是否在已选数组中，这个函数在模板中被用于控制复选框的选中样式。使用 wxs 处理这类纯模板层的判断逻辑，避免了在 JS 中对每条数据预处理添加标记字段，既减少了 setData 的数据量，又让模板代码更加清晰。"

t6g = "审批对话框采用模态弹窗设计。背景遮罩 .dialog-mask 使用 position: fixed 覆盖全屏，background: rgba(0,0,0,0.5) 半透明黑色，点击遮罩触发关闭。对话框本身 .dialog-box 是一个白色圆角卡片，使用 catchtap='stopPropagation' 阻止事件冒泡——catchtap 与 bindtap 的区别在于前者会阻止事件向上冒泡，因此点击对话框内部区域不会触发遮罩的关闭事件。对话框分标题栏、审批意见输入区和底部按钮三部分。"

# 七、AI对话
t7a = "AI 智能对话功能通过接入 DeepSeek 大语言模型，为师生提供了一个嵌入业务流程的智能助手。整个功能从前端的悬浮面板自定义组件、到工具层的 DeepSeek API 封装、再到跨页面的全局状态同步，构成了完整的技术链路。"

t7b = "前端在 utils/ai.js 中封装了独立的 AI 调用模块，使用 wx.request 向 DeepSeek 的 chat/completions 端点发送 POST 请求。请求体包含四个核心参数：model 指定为 deepseek-chat；messages 是符合 OpenAI 格式的对话消息数组，每条消息包含 role（system/user/assistant）和 content 字段；temperature 温度参数设为 0.7——这是一个 0 到 2 之间的值，控制输出的随机性，0.7 在创造性和确定性之间取得平衡；max_tokens 最大输出长度设为 2000，既保证回复完整又控制响应时长。模块暴露的 chat 方法支持多轮对话接收完整消息历史，ask 方法是对 chat 的简化封装用于单轮问答。"

t7c = "AI 对话的交互界面封装为一个名为 floating-panel 的自定义组件。组件使用 Component() 构造器定义，包含 properties 外部属性（如 title 面板标题）、data 内部数据（如 messages 对话消息列表、inputValue 输入框内容、isOpen 面板展开状态）和 methods 方法。组件在默认状态下以一个小圆形浮动按钮定位在页面右下角——使用 position: fixed、bottom: 120rpx、right: 40rpx 定位，width 和 height 均为 88rpx 圆形，border-radius: 50% 实现正圆，background 使用品牌蓝色，box-shadow 添加投影使其浮起于页面内容之上，z-index: 999 确保在大多数内容之上。"

t7d = "点击浮动按钮后面板展开。展开过程使用 CSS transition 过渡动画——面板从容器的 translateY(100%) 位移到 translateY(0)，配合 opacity 从 0 到 1，动画时长约 0.3s 使用 ease-out 缓出函数，产生从底部滑入的流畅效果。展开后的面板占据屏幕下半部分，顶部是标题栏。中间是 scroll-view 可滚动对话区域，消息气泡采用经典的聊天布局——用户消息靠右显示，使用蓝色背景白色文字的圆角气泡；AI 回复靠左显示，使用浅灰背景深色文字的圆角气泡。气泡内部文字使用 word-break: break-all 处理长文本换行。底部是固定在面板下沿的输入区域——input 输入框和发送按钮水平 flex 排列。"

t7e = "AI 对话的跨页面连续性通过全局状态管理实现。用户在个人中心点击 AI 对话菜单后，系统将 app.globalData.aiDialogEnabled 设为 true。当用户切换到其他 Tab 页面时，新页面的 onShow 生命周期函数检测到标记为 true，自动获取当前页的 floating-panel 实例并显示。当 AI 面板全屏展开时，页面通过 triggerEvent 向父页面发射自定义事件，父页面收到事件后调用 getTabBar().setData({ tabBarHidden: true }) 隐藏底部导航，避免遮挡对话界面。"

# 八、数据图表
t8a = "数据可视化看板是教师专属的分析工具，页面加载时通过 Promise.all 并行发起四个 API 请求——分别获取性别分布、年龄分布、生日月份分布和近 15 天请假趋势的数据。使用 Promise.all 而非 await 串行等待的设计使得四个请求几乎同时发出，总耗时约等于最慢的那个请求，而非四个请求的累加。请求全部完成后，通过解构赋值同时拿到四个结果，一次性设置数据并触发图表渲染，避免了数据到达时间不一致导致的渲染闪烁。"

t8b = "ECharts 是百度开源的数据可视化库，但其设计运行环境是 Web 浏览器的 DOM API，而微信小程序没有 DOM 概念。社区提供的 ec-canvas 组件解决了这个适配问题。ec-canvas 本质上是一个微信小程序自定义组件，其核心机制是在组件的 attached 生命周期中通过 wx.createSelectorQuery 获取新版 Canvas 2D 的节点对象，然后在 onInit 回调中将该节点传递给 ECharts 的 init 方法完成初始化。每个图表通过 ec 属性传入配置——ec 是一个包含 onInit 函数的对象。"

t8c = "以性别饼图为例，其 ECharts 配置包含：tooltip 提示框的 trigger 设为 'item'，表示悬浮在扇区上时显示详细数据浮层；legend 图例置于底部居中；series 中 type 设为 'pie'，radius 设为 ['45%','70%'] 数组——第一个值是内半径、第二个是外半径，两者不同即生成环形饼图（甜甜圈图），视觉上比实心饼图更现代；center 设为 ['50%','45%'] 将圆心略微上移为图例预留空间；label 的 formatter 设为 '{b}\n{d}%' 使每个扇区上显示名称和百分比两行文字。柱状图的配置使用笛卡尔坐标系——xAxis type 为 category 类目轴、yAxis type 为 value 数值轴，series type 为 bar。折线图的 series type 为 line，smooth: true 启用贝塞尔曲线平滑，areaStyle 设置半透明填充区域形成面积图效果。"

t8d = "考虑到 ECharts 压缩后约 500KB 的体积可能影响小程序加载速度，我们实现了一套完整的 Canvas 2D 手绘降级方案。所有 Canvas 操作在 JS 中通过 wx.createSelectorQuery 获取节点后执行。绘制流程的第一步是适配设备像素比——canvas.width 和 canvas.height 设为逻辑宽高乘以 dpr（通过 wx.getSystemInfoSync().pixelRatio 获取），然后 ctx.scale(dpr, dpr) 缩放绘图上下文。这个操作保证了在 Retina 高清屏上图形和文字不会模糊。"

t8e = "饼图的 Canvas 绘制使用 ctx.arc 方法绘制扇形。从 12 点钟方向（-Math.PI/2 弧度）开始，每个数据项的扇形角度为 item.value/total * 2 * Math.PI。绘制每个扇区前调用 ctx.beginPath() 开始新路径，ctx.moveTo(cx, cy) 移动到圆心，ctx.arc(cx, cy, radius, startAngle, startAngle+sliceAngle) 绘制弧线，ctx.closePath() 闭合路径，ctx.fill() 填充颜色。标签文字在扇形中间角度的延长线上绘制——标签离圆心距离根据百分比动态调整。柱状图使用 ctx.fillRect(x, y, barWidth, barHeight) 绘制矩形。折线图使用 ctx.moveTo 和 ctx.lineTo 连接数据点形成路径，ctx.stroke 描边绿色线条，再通过 ctx.fill 填充下方区域形成面积图。"

t8f = "在 WXML 模板中，ECharts 和 Canvas 降级通过 wx:if/wx:else 条件渲染实现互斥。useECharts 布尔变量在页面 onLoad 时通过 try/catch 包裹的 require 语句确定——成功加载则为 true，渲染 ec-canvas 组件；失败则为 false，渲染原生 canvas 标签（type 必须设为 '2d' 指定使用新版 Canvas 2D API）。图表布局使用 .charts-grid 网格容器——通过 flex-wrap: wrap 允许换行，每个 .chart-card 设置 flex-basis 控制宽度。卡片使用白色背景、24rpx 圆角和阴影。"

# 九、样式体系总结
t9a = "整个小程序的样式设计遵循了一套系统化的方法论。在设计变量层面，通过 CSS 自定义属性在 app.wxss 的 page 选择器上定义了一套完整的设计令牌，覆盖颜色、圆角、间距、字号、阴影五大类别。所有页面和组件的样式通过 var() 函数引用这些变量，实现了全局主题的一致性和可维护性。"

t9b = "在组件样式层面，按钮系统通过 .btn 基类定义通用样式（inline-flex 布局、圆角、字号、过渡动画），通过修饰类覆盖颜色、尺寸和宽度。标签系统采用同样的模式——.tag 基类加五个颜色变体。表单系统统一了 input、textarea、picker 三种控件的视觉风格——相同的高度 80rpx、相同的圆角 8rpx、相同的边框色和聚焦态蓝色高亮。卡片系统通过 .card 基类和 .form-card/.section-card 变体覆盖了不同场景的容器需求。"

t9c = "在布局层面，系统大量使用 flexbox 弹性布局——导航栏的左右分布、卡片内的信息排列、操作按钮的水平排列、地图图例的横向列表等。工具类 .flex/.flex-center/.flex-between 将高频的 flex 布局模式封装为单一样式类，在 WXML 中通过 class 属性直接引用，无需在每个组件中重复声明。"

t9d = "在交互状态层面，按钮通过 :active 伪类设置 opacity: 0.8 实现按下反馈；输入框通过 :focus 伪类切换边框颜色为品牌蓝提供聚焦反馈；禁用态通过 [disabled] 属性选择器设置 opacity: 0.7；消息提示通过 @keyframes 定义 slideDown 动画实现从顶部滑入的效果；对话框遮罩通过 rgba 半透明黑色和固定定位实现模态效果。这些交互细节遍布整个系统，共同构建了流畅、专业、有质感的用户体验。"

# 十、总结
t10a = "NETMAXXX 学生管理系统是一个从前端视觉到后端逻辑都经过精心设计的完整项目。前端方面，它充分利用了微信小程序的原生能力——map 组件实现地图可视化、wx.getLocation 实现 GPS 定位、wx.scanCode 实现摄像头扫码、wx.chooseMessageFile 实现文件选择、自定义 TabBar 实现角色感知的导航、CSS 变量实现主题统一管理、Canvas 2D 实现图表降级渲染、自定义组件实现 AI 对话的封装复用。后端方面，Django ORM 的 Q 查询和聚合统计、JWT 双 Token 的并发刷新机制、openpyxl 的 Excel 智能解析、bcrypt 的密码安全存储、基于角色的接口级权限控制，共同构成了稳健的服务端架构。"

t10b = "四个核心功能模块各具技术特色：签到模块深度融合了微信的定位和扫码硬件能力，通过 map 组件的 markers/circles/callout/label 属性系统实现了丰富的地图交互；请假模块实现了从申请到审批的完整工作流，利用 wxs 视图层脚本优化了模板渲染性能；AI 对话通过悬浮面板自定义组件和 DeepSeek API 集成，实现了跨页面的智能助手；数据图表通过 ECharts 优先加 Canvas 2D 降级的双引擎方案，在效果和兼容性之间取得了工程上的平衡。整个系统的样式设计遵循了设计令牌、原子化工具类、组件样式三层架构，保证了视觉的一致性和代码的可维护性。感谢大家的聆听，欢迎提问。"

# ========== 代码片段 ==========
code_pwd = [
    "// WXML — 密码显示切换控件",
    '<view class="pwd-toggle" bindtap="toggleLoginPassword">',
    '  <view class="pwd-checkbox {{loginShowPassword ? \'checked\' : \'\'}}">',
    '    <text wx:if="{{loginShowPassword}}">✓</text>',
    '  </view>',
    '  <text class="pwd-toggle-text">显示密码</text>',
    '</view>',
]

code_map = [
    "// map 组件 — 学生端全屏地图完整配置",
    '<map id="studentLocationMap" class="location-map"',
    '  latitude="{{studentMapLat}}" longitude="{{studentMapLng}}" scale="{{17}}"',
    '  markers="{{studentMapMarkers}}" circles="{{studentMapCircles}}" show-location="{{false}}" />',
    "",
    "// markers 数组构建 — 红色我的位置 + 蓝色教师位置",
    'markers.push({',
    '  id: 0, latitude: studentLat, longitude: studentLng, width: 32, height: 32,',
    "  callout: { content: '我的位置', color: '#E74C3C', display: 'ALWAYS', borderRadius: 8, padding: 6 },",
    "  label: { content: '我', color: '#E74C3C', fontSize: 12, anchorX: 0, anchorY: -12 }",
    '});',
    "// circles 覆盖物 — 半透明蓝色签到范围圈",
    'circles.push({',
    '  latitude: teacherLat, longitude: teacherLng, radius: range,',
    "  color: '#1677FF66', fillColor: '#1677FF1A', strokeWidth: 2",
    '});',
]

code_wxs = [
    "// wxs 模块 — 视图层工具函数（运行在渲染层，不经过 JS 线程）",
    "// leave-manage.wxs",
    'module.exports = {',
    '  isSelected: function(arr, id) {',
    '    return arr.some(function(item) { return item.id === id; });',
    '  }',
    '};',
]

code_canvas = [
    "// Canvas 2D 降级方案 — 饼图绘制核心算法",
    'let startAngle = -Math.PI / 2;',
    'data.forEach((item, i) => {',
    '  const sliceAngle = (item.value / total) * 2 * Math.PI;',
    '  ctx.beginPath(); ctx.moveTo(cx, cy);',
    '  ctx.arc(cx, cy, radius, startAngle, startAngle + sliceAngle);',
    '  ctx.closePath();',
    '  ctx.fillStyle = colors[i % colors.length]; ctx.fill();',
    '  // 在扇区中间方向绘制标签',
    '  const midAngle = startAngle + sliceAngle / 2;',
    '  const labelX = cx + Math.cos(midAngle) * (radius + 20);',
    '  const labelY = cy + Math.sin(midAngle) * (radius + 20);',
    "  ctx.fillText(item.name + ':' + item.value, labelX, labelY);",
    '  startAngle += sliceAngle;',
    '});',
]

# ========== 组装文档 ==========
# 封面
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(120)
r = p.add_run(t_cover_title); SF(r, '黑体', 'Arial', 26, True, RGBColor(0x16,0x77,0xFF))
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run(t_cover_sub); SF(r2, '宋体', 'Consolas', 16, color=RGBColor(0x66,0x66,0x66))
doc.add_paragraph()

# 一
T("一、项目概述与技术架构", 1)
for t in [t1a, t1b, t1c, t1d, t1e]: P(t)

# 二
T("二、登录与注册页 — 视觉设计与交互实现", 1)
for t in [t2a, t2b, t2c, t2d, t2e]: P(t)
T("2.1 密码显示切换", 2)
for line in code_pwd: C(line)

# 三
T("三、角色权限与自定义 TabBar", 1)
for t in [t3a, t3b, t3c, t3d]: P(t)

# 四
T("四、学生信息管理", 1)
for t in [t4a, t4b, t4c]: P(t)

# 五
T("五、签到功能（重点详解）", 1)
P(t5_intro)
T("5.1 教师发布签到 — 表单控件与定位权限", 2)
for t in [t5a, t5b]: P(t)
T("5.2 学生扫码签到 — 摄像头调用", 2)
P(t5c)
T("5.3 学生位置签到 — map 原生组件深度应用", 2)
for t in [t5d, t5e, t5f]: P(t)
for line in code_map: C(line)
T("5.4 教师考勤管理 — 地图+列表双视图", 2)
for t in [t5g, t5h]: P(t)

# 六
T("六、请假功能（重点详解）", 1)
for t in [t6a, t6b, t6c, t6d, t6e]: P(t)
T("6.1 教师端 — wxs 视图层脚本与审批对话框", 2)
for t in [t6f, t6g]: P(t)
for line in code_wxs: C(line)

# 七
T("七、AI 对话功能（重点详解）", 1)
for t in [t7a, t7b]: P(t)
T("7.1 floating-panel 悬浮面板组件", 2)
for t in [t7c, t7d]: P(t)
T("7.2 跨页面状态同步", 2)
P(t7e)

# 八
T("八、数据图表展示（重点详解）", 1)
for t in [t8a, t8b, t8c]: P(t)
T("8.1 Canvas 2D 降级方案", 2)
for t in [t8d, t8e, t8f]: P(t)
for line in code_canvas: C(line)

# 九
T("九、前端样式体系总结", 1)
for t in [t9a, t9b, t9c, t9d]: P(t)

# 十
T("十、总结", 1)
for t in [t10a, t10b]: P(t)

# 保存
out = r'N:\课程备份\小程序\大作业\小程序演示讲解稿_v3.docx'
doc.save(out)
print(f'[OK] {out}')
