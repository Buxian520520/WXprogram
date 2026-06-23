"""
v4 讲解稿 — 前后端完整版。每个模块：前端实现 → 后端实现。
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

def SF(run, cn, en, size, bold=False, color=None):
    run.font.size = Pt(size); run.bold = bold; run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rF = rPr.find(qn('w:rFonts'))
    if rF is None: rF = OxmlElement('w:rFonts'); rPr.insert(0, rF)
    rF.set(qn('w:eastAsia'), cn); rF.set(qn('w:ascii'), en); rF.set(qn('w:hAnsi'), en)
    if color: run.font.color.rgb = color

def T(text, level=1):
    h = doc.add_heading(text, level=level); h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in h.runs: SF(r, '黑体', 'Arial', {0:22,1:16,2:14}.get(level,14), True)

def P(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); SF(r, '宋体', 'Consolas', 12)

def C(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74); p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); SF(r, '宋体', 'Consolas', 10.5, color=RGBColor(0x16,0x77,0xFF))

# ========== 封面 ==========
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(120)
r = p.add_run("NETMAXXX 学生管理系统"); SF(r, '黑体', 'Arial', 26, True, RGBColor(0x16,0x77,0xFF))
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("微信小程序演示讲解稿（前后端完整版 v4）"); SF(r2, '宋体', 'Consolas', 16, color=RGBColor(0x66,0x66,0x66))
doc.add_paragraph()

# ================================================================
# 一、项目概述
# ================================================================
T("一、项目概述与技术架构", 1)

P("大家好，我今天展示的项目是 NETMAXXX 学生管理系统，一个基于微信小程序平台开发、Django 框架支撑后端的综合校园信息管理系统。系统支持学生和教师两种角色，涵盖学生档案管理、课堂签到、请假审批、数据统计分析和 AI 智能对话等功能。下面我将按照功能模块逐一介绍前后端的完整技术实现。")

T("1.1 前端技术栈", 2)

P("前端采用微信原生框架开发，每个页面由四种文件组成：WXML 模板文件负责界面结构，使用微信自己的组件体系如 view、text、image、input、picker、map、slider、scroll-view、textarea 等；WXSS 样式文件负责视觉效果，语法兼容 CSS，支持 flexbox 弹性布局、CSS 变量（自定义属性）、关键帧动画、过渡动画等现代特性；JS 逻辑文件使用 Page() 或 Component() 构造器注册，通过 setData 方法驱动视图更新，调用 wx.xxx 系列 API 使用微信硬件能力；JSON 配置文件负责页面级设置。")

P("在样式架构上，全局样式文件 app.wxss 在 page 选择器上通过 CSS 自定义属性定义了一套设计令牌——包括 --color-primary（品牌蓝 #1677FF）、--color-success（成功绿 #00B42A）、--color-warning（警告橙 #FF7D00）、--color-danger（危险红 #F53F3F）等语义化颜色；--radius-sm/md/lg 三级圆角变量；--spacing-xs/sm/md/lg 四级间距变量；--font-xs/sm/md/lg/xl 五级字号变量。在此基础上还定义了 .btn 按钮体系（基类 + .btn-primary/.btn-success/.btn-danger/.btn-outline/.btn-sm/.btn-block 修饰类）、.tag 标签体系（基类 + 五色变体）、.card 卡片体系、flex 布局工具类等原子化 CSS 类。")

T("1.2 后端技术栈", 2)

P("后端基于 Django 5.2 框架构建，使用 MySQL 数据库。Django 项目结构分为：StudentV4BE 项目配置目录（settings.py 数据库/CORS/中间件配置、urls.py 根路由）；apps/student 应用目录（models.py 数据模型、views.py 学生管理视图、auth_views.py JWT 认证视图、leave_views.py 请假管理视图、teacher_views.py 教师管理视图）。前后端通过 RESTful API 通信——前端使用封装的 api.js 模块发起 HTTP 请求，后端视图函数解析 JSON 请求体并返回 JsonResponse。认证使用 JWT 双 Token 机制（Access Token 30 分钟 + Refresh Token 7 天），密码使用 bcrypt 哈希存储。后端通过 natapp 内网穿透将本地 8000 端口映射到公网域名。")

C("// Django settings.py — 数据库与CORS配置")
C("DATABASES = { 'default': { 'ENGINE': 'django.db.backends.mysql',")
C("    'NAME': 'studentv4db', 'USER': 'root', 'HOST': 'localhost', 'PORT': '3306' }}")
C("CORS_ALLOWED_ORIGINS = ['http://localhost:63342', 'http://127.0.0.1:5500']")

# ================================================================
# 二、认证与权限体系
# ================================================================
T("二、认证与权限体系", 1)

T("2.1 前端 — 登录注册页", 2)

P("登录页采用全屏渐变背景——CSS 的 linear-gradient(135deg, #1677FF, #69b1ff)，配合 min-height: 100vh 覆盖全屏。内容由 Logo 标题区、模式切换选项卡和表单卡片三部分组成。模式切换使用半透明磨砂玻璃风格——background: rgba(255,255,255,0.2)、border-radius: 40rpx 胶囊形状，选中态通过 .mode-tab-active 切换为白色背景加品牌蓝文字，transition: all 0.3s 平滑过渡。表单卡片使用白色背景、border-radius: 24rpx 大圆角、box-shadow: 0 8rpx 32rpx rgba(0,0,0,0.1) 深阴影从背景浮起。输入框聚焦时通过 :focus 伪类从浅灰底切换为白底蓝边框。消息提示通过 @keyframes slideDown 动画从顶部滑入。")

P("密码框下方增加了显示/隐藏切换控件——一个 36rpx 自定义复选框 + 文字标签，点击通过 bindtap 翻转 showPassword 布尔值，input 组件的 password 属性绑定 {{!showPassword}} 表达式——true 时明文、false 时圆点。注册表单根据角色动态展示不同字段——学生需填写姓名/性别/生日/手机/邮箱/地址，教师只需填写工号和姓名，通过 wx:if 条件渲染切换。")

T("2.2 后端 — JWT 认证与注册", 2)

P("后端的认证逻辑在 auth_views.py 中实现，包含 login、register 和 refresh_token_view 三个视图函数。JWT Token 使用 PyJWT 库生成，算法为 HS256，密钥为 secret-key。Access Token 的 payload 包含 username、role、display_name、type='access'、exp（30 分钟过期）和 iat（签发时间）；Refresh Token 结构类似但 type='refresh' 且过期时间为 7 天。登录时先通过 User.objects.get(username=username) 查找用户，再使用 bcrypt.checkpw 验证密码哈希，成功后调用 generate_tokens 生成双 Token 返回。")

P("注册接口同时完成三件事：使用 bcrypt.hashpw 加密密码；在 Student 表或 Teacher 表中创建/更新对应的业务档案（通过 update_or_create 避免重复）；在 users 表中创建认证记录，通过 user_id 字段关联到业务表（学生关联 sno、教师关联 Teacher.id）。学生注册时自动将 username 规范化为学号数字，使用正则验证 95xxx 格式；教师注册时自动以工号作为 username。注册成功后直接返回 Token，用户无需再次登录。")

C("// auth_views.py — JWT Token 生成")
C("payload = { 'username': username, 'role': role, 'display_name': display_name,")
C("    'type': 'access', 'exp': datetime.utcnow() + timedelta(minutes=30), 'iat': datetime.utcnow() }")
C("access_token = jwt.encode(payload, SECRET_KEY, algorithm='HS256')")

T("2.3 前端 — 角色感知的 TabBar 与路由", 2)

P("系统使用微信小程序的自定义 TabBar 机制。组件定义在 custom-tab-bar 目录下，使用 Component() 构造器，在 attached 生命周期中调用 updateTabs 方法——读取 app.globalData.userRole 从预定义数组中过滤当前角色可见的标签。教师端显示学生信息/数据中心/签到/我的四个标签，学生端显示签到/请假/我的三个标签。切换通过 wx.switchTab API。TabBar 在签到地图或 AI 全屏时通过 getTabBar().setData({tabBarHidden:true}) 隐藏。")

T("2.4 后端 — 接口级角色权限控制", 2)

P("权限控制的核心在后端视图函数层面。以 get_students 为例——函数先从 HTTP_AUTHORIZATION 头解析 Bearer Token，调用 verify_token 验证并提取 role 字段。如果 role 为 student，则只通过 Student.objects.filter(sno=user.user_id) 返回该学生自己的记录；如果 role 为 teacher，则返回 Student.objects.all() 全量数据。类似的，请假管理中 leave_views.py 的 get_leaves 函数学生只能看到自己的请假记录，教师可看到全部并可筛选；教师专属接口（如 update_teacher_info）会先校验 role == 'teacher'，否则返回 403。这种在视图函数层面的角色分支，保证了即使前端被绕过，后端也能拦住越权请求。")

C("// views.py — 角色数据隔离")
C("if user_info['role'] == 'student':")
C("    student = Student.objects.filter(sno=user.user_id).first()")
C("    students = [student] if student else []")
C("else:")
C("    students = Student.objects.all().values()")

# ================================================================
# 三、学生信息管理
# ================================================================
T("三、学生信息管理", 1)

T("3.1 前端 — 列表页与表单页", 2)

P("学生列表页顶部搜索栏使用 flex 水平布局——input 设置 flex:1 占据剩余空间。学生卡片使用水平 flex 布局——左侧 image 组件 mode='aspectFill' 展示头像，无头像时显示取首字的蓝色圆形占位符；中间展示姓名/性别标签/学号/生日，性别通过条件类名蓝男红女；右侧操作按钮仅在教师身份时 wx:if 渲染。复选框选中时通过 .checkbox-checked 类显示蓝底白勾。分页控件使用 picker 选择每页条数（5/10/50/100），配合上一页/下一页按钮。批量选择支持当前页全选/取消全选——通过 Set 去重合并实现。")

P("表单页支持 add/edit/view 三种模式，通过 URL 参数 mode 区分。学号输入框在非编辑模式下绑定 bindblur 事件调用 sno/check/ 接口做异步唯一性校验。头像上传使用 wx.chooseImage 选择图片后调用 wx.uploadFile 上传，成功后将返回的文件名保存到 form.image。表单提交前通过 validator.js 做正则校验——学号 /^[9][5]\\d{3}$/、姓名 /^[一-龥]{2,5}$/、手机 /^[1][35789]\\d{9}$/。")

T("3.2 后端 — Django ORM 数据层", 2)

P("Student 模型使用 models.IntegerField 定义 sno 为主键（primary_key=True），Gender 使用 choices 参数限制男/女。Meta 类中 db_table='Student' 指定表名。后端提供完整的 CRUD 接口——add_student 创建记录并同步在 users 表中创建登录账号（默认密码 123456 用 bcrypt 加密）；update_student 通过 Student.objects.get(sno=data['sno']) 获取后逐字段修改；delete_student 删除学生记录的同时清理 users 表中关联账号。所有操作完成后返回全量学生列表供前端刷新。")

P("搜索功能使用 Django Q 对象构建 OR 条件查询——Q(sno__icontains=inputstr) | Q(name__icontains=inputstr) | Q(gender__icontains=inputstr) | Q(mobile__icontains=inputstr) | Q(email__icontains=inputstr) | Q(address__icontains=inputstr)。六个字段的 icontains 模糊匹配意味着输入任意关键词都能命中。")

T("3.3 Excel 导入导出", 2)

P("前端导入：wx.chooseMessageFile({count:1, type:'file', extension:['xlsx','xls']}) 选择文件，api.uploadExcel 通过 wx.uploadFile 上传。导出：wx.downloadFile 下载后端生成的 Excel，通过 wx.shareFileMessage 分享或 wx.openDocument 打开。后端使用 openpyxl 库——read_excel_dict 函数智能识别表头：先读第一行判断是否包含学号/姓名等关键字，有表头则按列名映射到字段名，无表头按固定列序解析。导入采用 update_or_create 方法，已存在的学号更新资料、不存在的创建新记录，同时自动创建 users 账号。容错机制：单行异常不中断流程，记录错误后继续处理后续行，最后返回成功数/失败数/失败学号列表。")

C("// views.py — Excel导入核心逻辑")
C("Student.objects.update_or_create(sno=sno, defaults={")
C("    'name': name, 'gender': gender, 'birthday': birthday,")
C("    'mobile': mobile, 'email': email, 'address': address })")

# ================================================================
# 四、签到功能
# ================================================================
T("四、签到功能（重点详解）", 1)

P("签到功能深度融合了微信小程序的 GPS 定位和摄像头扫码能力。整个模块在单一 checkin.js 页面中通过 wx:if 条件渲染实现了教师和学生两套 UI。下面从前端和后端两个维度详细介绍。")

T("4.1 前端 — 教师发布签到", 2)

P("教师发布表单使用三种原生组件：input（bindinput 实时同步标题）、picker mode='selector'（选择二维码/位置签到）、slider（min=10 max=2000 step=10 activeColor='#1677FF' block-size=20 调整签到范围）。位置签到时调用 _getLocation 获取 GPS——先用 wx.getSetting 检查 scope.userLocation 的三层权限状态（true 已授权/false 曾拒绝/undefined 未询问），已拒绝时通过 wx.showModal + wx.openSetting 引导用户去系统设置开启；已授权或未询问时调用 wx.getLocation({type:'gcj02', isHighAccuracy:true, highAccuracyExpireTime:15000})。获取坐标后通过 api.post('checkin/create/', {type, title, range, lat, lng}) 发布任务。")

T("4.2 前端 — 学生扫码签到", 2)

P("学生通过 wx.scanCode({onlyFromCamera:true}) 调用摄像头扫码——onlyFromCamera 禁止从相册选取，保证课堂物理场景。扫码成功后从 res.result 获取签到的码，连同 task_id 通过 api.post('checkin/sign/', {task_id, type:'qrcode', code}) 提交。也可手动输入签到的码后点击签到按钮提交。签到成功或失败都有 toast 轻提示反馈。")

T("4.3 前端 — 学生位置签到与 map 地图组件", 2)

P("位置签到是技术实现最复杂的环节，核心依赖微信原生 map 组件（底层为腾讯地图）。学生点击获取位置并签到后，系统先获取 GPS 坐标，然后打开全屏地图界面。map 组件配置：latitude/longitude 设置中心点（优先教师坐标）、scale=17（约 50 米比例尺）、markers 数组包含两个标记点——红色 32x32 标记我的位置（callout 气泡 display='ALWAYS' 始终显示、borderRadius=8 圆角、padding=6 内边距），蓝色 28x28 标记教师位置。circles 数组定义半透明签到范围圈——以教师坐标为圆心、range 为半径，color='#1677FF66' 边框半透明蓝、fillColor='#1677FF1A' 填充更透明。")

P("地图界面布局：外层 .location-map-overlay 使用 position:fixed 全屏覆盖、z-index 高层级。图例栏 .map-legend 使用 position:absolute 左上角定位、半透明白底圆角阴影。底部操作面板 .location-panel 使用白色背景 + border-radius: 32rpx 32rpx 0 0 大顶部圆角模拟底部弹出卡片效果，包含取消（.btn-outline）和确认签到（.btn-primary）两个按钮。确认签到调用 api.post('checkin/sign/', {task_id, type:'location', lat, lng})。")

C("// map 组件 — markers 与 circles 配置")
C("markers.push({ id:0, latitude:studentLat, longitude:studentLng, width:32, height:32,")
C("  callout:{content:'我的位置',color:'#E74C3C',display:'ALWAYS',borderRadius:8,padding:6},")
C("  label:{content:'我',color:'#E74C3C',fontSize:12,anchorX:0,anchorY:-12} });")
C("circles.push({ latitude:teacherLat, longitude:teacherLng, radius:range,")
C("  color:'#1677FF66', fillColor:'#1677FF1A', strokeWidth:2 });")

T("4.4 前端 — 教师考勤管理", 2)

P("教师可查看签到记录弹窗（.dialog-mask 半透明遮罩 + .dialog-box 白色圆角卡片，catchtap 阻止冒泡），也可打开全屏地图查看所有学生签到位置——map 使用相同配置但 markers 数量更多，学生标记用 24x24 较小尺寸，callout display='BYCLICK' 点击显示，点击标记触发 bindmarkertap 事件通过 e.detail.markerId 定位学生并 toast 显示信息。")

P("考勤详情面板展示完整出勤统计——标题区水平 flex 排列已签人数（绿色）/未签人数（红色）/总人数（灰色）。列表区 scroll-view 设置 max-height:60vh 可滚动，每行显示姓名学号/状态标签/操作按钮。未签到学生显示绿色补签按钮调用 checkin/record/add/，已签到学生显示红色删除按钮调用 checkin/record/remove/。补签设计解决了学生手机故障等特殊情况。")

T("4.5 后端 — 签到 API 设计", 2)

P("虽然当前后端 urls.py 中尚未注册签到相关路由（前端 checkin.js 调用的接口标记为待迁移），但从前端代码可以清晰推导出后端 API 的设计规范。checkin/create/ 接口接收 POST 请求，body 包含 type（qrcode/location）、title、range、lat、lng，后端创建 CheckinTask 记录并返回任务信息。checkin/sign/ 接口接收 task_id、type、code（二维码）或 lat/lng（位置），后端创建 CheckinRecord 记录，需校验同一 task 下 sno 唯一（uk_task_sno 约束）。checkin/end/ 接口将任务 status 从 active 改为 ended 并记录 ended_at。checkin/attendance/ 接口需关联 Student 全表与 CheckinRecord 做 LEFT JOIN 比对，区分已签和未签学生。checkin/history/ 返回 status='ended' 的历史任务列表。checkin/delete/ 级联删除任务及其所有签到记录。")

# ================================================================
# 五、请假功能
# ================================================================
T("五、请假功能（重点详解）", 1)

T("5.1 前端 — 学生端请假表单", 2)

P("学生端请假页采用卡片折叠式设计。顶部表单卡片 .form-card——白色背景、24rpx 圆角、32rpx 内边距、轻微阴影。卡片头部水平 flex，右侧切换按钮通过三元表达式动态显示收起或展开。表单区域使用 wx:if（而非 hidden）控制显示——wx:if 完全销毁/重建 DOM，收起时清空所有输入状态，下次展开获得全新表单。请假类型使用 picker mode='selector'。开始/结束日期使用 picker mode='date'——关键设计是结束日期选择器设置了 start 属性绑定开始日期值，原生日期选择器自动置灰早于 start 的日期，从 UI 层杜绝日期倒挂。请假原因使用 textarea（maxlength=500），右下角绝对定位字符计数器实时显示 {{leaveForm.reason.length}}/500。表单提交前通过 validator.validateLeaveForm 校验——类型/日期/原因均不能为空，结束日期不能早于开始，原因至少 5 字符。error 对象中对应字段信息通过 wx:if 条件渲染在字段下方显示红色错误提示。")

P("请假记录列表采用卡片式布局。每条 .leave-item 顶部水平展示两个 tag 标签——请假类型（病假橙色 tag-warning/事假灰色 tag-info）和审批状态（待审批橙色/已批准绿色 tag-success/已拒绝红色 tag-danger），全部复用全局 .tag 样式体系。状态筛选器使用 picker，选择后通过 JS filter 方法本地过滤——statusIndex 为 0 不过滤，1~3 分别匹配 pending/approved/rejected。")

T("5.2 前端 — 教师端审批工作流", 2)

P("教师端使用 wxs（WeiXin Script）视图层脚本——wxs 运行在渲染层而非 JS 线程，通过 <wxs module='utils' src='./leave-manage.wxs'></wxs> 引入。导出的 isSelected 函数在模板中直接调用控制复选框选中样式，避免在 JS 中对每条数据预处理，减少 setData 数据量。")

P("审批对话框采用模态设计——.dialog-mask 遮罩（position:fixed + rgba(0,0,0,0.5)），.dialog-box 白色圆角卡片使用 catchtap='stopPropagation' 阻止冒泡（catchtap 与 bindtap 区别：前者阻止事件向上传播，点击卡片内部不会触发遮罩关闭）。标题根据 approvalType 动态显示批准/拒绝。审批意见 textarea 可选填。提交时调用 leave/approve/ 或 leave/reject/ 接口，传 leaveId 和 comment。")

C("// wxs — 视图层脚本（运行在渲染层）")
C("module.exports = {")
C("  isSelected: function(arr, id) {")
C("    return arr.some(function(item) { return item.id === id; });")
C("  }")
C("};")

T("5.3 后端 — Leave 数据模型", 2)

P("Leave 模型定义在 apps/student/models.py 中，包含 12 个字段。核心字段：student_id（IntegerField，关联 Student.SNo）、student_name 和 student_no（冗余字段，避免教师端查询时 JOIN Student 表，提升查询性能）、leave_type（CharField，choices 限制为事假/病假/其他）、start_date 和 end_date（DateField）、days（IntegerField，提交时自动计算 = (end_date-start_date).days+1）、reason（TextField）、status（CharField，choices 限制为 pending/approved/rejected，default='pending'）、apply_time（DateTimeField，auto_now_add=True 自动记录创建时间）、teacher_id（审批教师 ID）、approval_comment（TextField，审批意见）、approval_time（DateTimeField，审批时间）。Meta 类 db_table='Leave' 指定表名。")

T("5.4 后端 — 请假管理 API", 2)

P("请假管理的视图函数在 leave_views.py 中，共 6 个接口。get_leaves（GET）：从 Token 解析用户角色——学生只返回 student_id 等于自己 sno 的记录，教师返回全部并可接收 student_name 和 status 两个查询参数做筛选，按 apply_time 倒序排列。apply_leave（POST）：学生提交申请，后端从 Token 获取用户信息，通过 Student 表查学生姓名和学号，计算请假天数，创建 Leave 记录。approve_leave / reject_leave（POST）：教师审批，先校验角色为 teacher，再通过 Leave.objects.get(id=leave_id) 获取记录，校验 status 必须为 pending（防止重复审批），更新 status/approval_comment/approval_time/teacher_id 四个字段。delete_leave（POST）：接收 leaveIds 数组，学生只能删除自己的且状态为 pending 的记录，教师可删除任意记录。")

C("// leave_views.py — 审批核心逻辑")
C("leave = Leave.objects.get(id=leave_id)")
C("if leave.status != 'pending':")
C("    return JsonResponse({'code':0, 'msg':'该请假记录已处理'})")
C("leave.status = 'approved'; leave.approval_comment = comment")
C("leave.approval_time = datetime.now(); leave.teacher_id = user.user_id")
C("leave.save()")

# ================================================================
# 六、AI 对话功能
# ================================================================
T("六、AI 对话功能（重点详解）", 1)

T("6.1 前端 — DeepSeek API 封装", 2)

P("前端在 utils/ai.js 中封装了独立的 AI 调用模块。使用 wx.request 向 DeepSeek 的 https://api.deepseek.com/v1/chat/completions 端点发送 POST 请求。请求体包含：model='deepseek-chat'、messages（OpenAI 格式消息数组，每条含 role: system/user/assistant 和 content）、temperature=0.7（控制随机性，0~2，值越高越多样）、max_tokens=2000（限制回复长度）。模块暴露 chat(messages) 支持多轮对话（传入完整历史）、ask(question) 简化单轮问答。API Key 内置默认值，也可通过 setApiKey 自定义。")

T("6.2 前端 — floating-panel 悬浮面板组件", 2)

P("AI 交互界面封装为 floating-panel 自定义组件，使用 Component() 构造器定义——properties 接收外部属性（如 title）、data 维护内部状态（messages 对话列表、inputValue 输入内容、isOpen 展开状态）、methods 提供操作方法。默认状态为右下角圆形浮动按钮——position:fixed、bottom:120rpx、right:40rpx、88rpx 圆形、border-radius:50%、品牌蓝背景、box-shadow 投影、z-index:999。点击展开——transition 过渡动画（translateY(100%)→translateY(0) + opacity 0→1，0.3s ease-out）。展开后为半屏面板：顶部标题栏、中间 scroll-view 对话列表（用户消息右侧蓝色气泡、AI 回复左侧灰底气泡，word-break:break-all 处理长文本）、底部固定输入区（input + 发送按钮）。")

T("6.3 前端 — 跨页面状态同步", 2)

P("用户在个人中心点击 AI 对话后，app.globalData.aiDialogEnabled 设为 true。其他页面在 onShow 生命周期中检测该标记，若为 true 则自动获取当前页 floating-panel 实例并显示触发按钮，实现跨 Tab 页面的 AI 对话连续性。面板全屏展开时通过 triggerEvent('statechange', {isOpen:true}) 向父页面发射事件，父页面调用 getTabBar().setData({tabBarHidden:true}) 隐藏底部导航避免遮挡。")

T("6.4 后端 — AI 服务说明", 2)

P("AI 对话功能直接调用 DeepSeek 的云端 API，不经过本项目的 Django 后端。这种设计将 AI 推理的计算压力完全放在云端，小程序端只负责发送请求和渲染回复，不需要额外的后端算力支持。DeepSeek API 兼容 OpenAI 的 Chat Completions 接口格式，返回的 choices[0].message.content 即为 AI 的文本回复。")

# ================================================================
# 七、数据图表展示
# ================================================================
T("七、数据图表展示（重点详解）", 1)

T("7.1 前端 — 双引擎渲染方案", 2)

P("图表页加载时使用 Promise.all 并行发起四个 API 请求——getSourceChartData（性别分布）、getAgeChartData（年龄分布）、getTypeChartData（生日月份）、getLeaveTrendData（请假趋势）。Promise.all 使四个请求同时发出，总耗时约等于最慢请求而非累加。完成后通过解构赋值同时拿到四个结果，一次性 setData 并渲染，避免数据到达不同步导致闪烁。")

P("首选引擎为 ECharts，通过 ec-canvas 自定义组件适配微信小程序。ec-canvas 在 attached 生命周期中通过 wx.createSelectorQuery().select('#canvas').fields({node:true,size:true}) 获取 Canvas 2D 节点，在 onInit 回调中将节点传给 echarts.init 完成初始化。每个图表的 ec 属性包含 onInit 函数，内部调用 chart.setOption 设置配置——以性别饼图为例：tooltip 提示框 trigger='item'；legend 图例置于底部居中；series type='pie'，radius=['45%','70%'] 生成环形饼图，center=['50%','45%'] 为图例留空间，label formatter='{b}\\n{d}%' 显示名称和百分比。柱状图使用笛卡尔坐标系——xAxis type='category' + yAxis type='value' + series type='bar'。折线图 series type='line'，smooth:true 贝塞尔平滑，areaStyle 半透明填充。")

P("降级方案使用 Canvas 2D API 手绘。先适配设备像素比——canvas.width/height = 逻辑尺寸 x dpr（wx.getSystemInfoSync().pixelRatio），ctx.scale(dpr,dpr) 保证高清屏不模糊。饼图用 ctx.arc(cx,cy,radius,startAngle,startAngle+sliceAngle) 绘制扇形，ctx.fillText 在扇区中间角度的延长线上绘制标签。柱状图用 ctx.fillRect(x,y,barW,barH) 绘制矩形。折线图用 ctx.moveTo+ctx.lineTo 连接数据点，ctx.fill 填充下方区域。两套方案通过 wx:if/wx:else 条件渲染互斥，useECharts 由 try/catch require 确定。图表卡片使用 flex-wrap 网格布局。")

C("// Canvas 2D 降级 — 饼图扇区绘制")
C("ctx.beginPath(); ctx.moveTo(cx, cy);")
C("ctx.arc(cx, cy, radius, startAngle, startAngle + sliceAngle);")
C("ctx.closePath(); ctx.fillStyle = colors[i]; ctx.fill();")

T("7.2 后端 — 图表数据统计", 2)

P("四个图表数据接口均在 views.py 中实现。update_Sex（性别分布）：使用 Student.objects.filter(gender='男').count() 和 filter(gender='女').count() 分别统计，返回 [{name:'男',value:N},{name:'女',value:N}]。get_ageData（年龄分布）：遍历所有学生，用当前年份减 birthday.year，再根据月日比较修正未到生日的情况，按年龄值分组计数，返回 [{name:'X岁',value:N}]。update_Birthday（生日月份）：使用 Django ORM 的 annotate(month=ExtractMonth('birthday')) 提取月份，.values('month').annotate(count=Count('sno')) 按月分组统计，确保 1~12 月都有数据（即使为 0），映射为一月~十二月中文名。get_leave_trend_data（请假趋势）：查询 Leave 表中 apply_time 在最近 15 天内的记录，使用 annotate(day=TruncDate('apply_time')) 按天截断，.values('day').annotate(count=Count('id')) 按天计数，对无记录日期补 0，返回 [{name:'YYYY-MM-DD',value:N}]。")

C("// views.py — 生日月份统计（Django ORM 聚合）")
C("birth_month_data = Student.objects.annotate(")
C("    month=ExtractMonth('birthday')")
C(").values('month').annotate(count=Count('sno')).order_by('month')")

# ================================================================
# 八、Token 管理与网络层
# ================================================================
T("八、Token 管理与网络层", 1)

T("8.1 前端 — api.js 请求封装", 2)

P("utils/api.js 是整个前端与后端通信的统一入口。模块封装了 request 核心方法——自动从 Storage 读取 Token 并注入 Authorization: Bearer xxx 请求头，默认 Content-Type: application/json。关键设计是 401 拦截与并发刷新控制：当响应状态码为 401 时，不立即跳转登录页，而是先检查 isRefreshing 标记。如果当前没有刷新进行中，则置标记为 true 并调用 doRefreshToken；如果已有刷新进行中，则将重试回调推入 refreshQueue 队列。刷新成功后遍历队列执行所有等待的回调，每个回调以携带新 Token 的方式重新发起原始请求。这个机制解决了多请求同时返回 401 导致重复刷新的并发问题。刷新失败则清除 Token 并 reLaunch 到登录页。")

P("便捷方法——get(url, params) 将 params 对象转为查询字符串拼接 URL；post(url, data) 直接传 JSON body；uploadFile(filePath, name) 和 uploadExcel(filePath) 封装 wx.uploadFile；login(username, password, role) 和 register(regData) 封装认证接口。Token 管理——saveTokens 将 access_token/refresh_token/expireTime（当前时间+29分钟安全边界）存入 Storage；getAccessToken 检查过期时间后返回有效 Token；clearTokens 清除所有登录凭证。")

C("// api.js — 401 并发刷新控制")
C("if (res.statusCode === 401 && !options._retry) {")
C("  if (isRefreshing) {")
C("    refreshQueue.push(() => request({...options, _retry:true})); return;")
C("  }")
C("  isRefreshing = true;")
C("  doRefreshToken().then(() => {")
C("    refreshQueue.forEach(cb => cb()); refreshQueue = []; doRequest();")
C("  });")
C("}")

T("8.2 后端 — JWT 刷新接口", 2)

P("refresh_token_view 接收 POST 请求中的 refresh_token，调用 verify_token(token, token_type='refresh') 校验——验证签名、检查 type 是否为 refresh、检查是否过期。校验通过后从 payload 取出 username 和 role，生成新的 access_token（不更新 refresh_token）。verify_token 函数内部使用 jwt.decode 解码，捕获 ExpiredSignatureError（过期）和 InvalidTokenError（无效）两种异常，返回 None 表示验证失败。")

# ================================================================
# 九、总结
# ================================================================
T("九、总结", 1)

P("NETMAXXX 学生管理系统是一个前后端完整、技术栈丰富的综合性项目。前端充分利用微信小程序原生能力——map 组件地图可视化（markers/circles/callout/label 属性体系）、wx.getLocation GPS 定位（gcj02 坐标系 + 高精度模式 + 三层权限处理）、wx.scanCode 摄像头扫码（onlyFromCamera 防作弊）、wx.chooseMessageFile/wx.downloadFile 文件操作、自定义 TabBar 角色感知导航、CSS 变量设计令牌体系（5 类变量 + 原子化工具类 + 组件样式三层架构）、Canvas 2D 手绘图表降级方案、floating-panel 自定义组件封装 AI 对话、wxs 视图层脚本优化模板性能。")

P("后端基于 Django 5.2 + MySQL，实现了完整的 RESTful API 体系——JWT 双 Token 认证（HS256 + bcrypt 密码哈希）、Django ORM 数据建模（Q 查询/聚合统计/ExtractMonth/TruncDate）、基于角色的接口级权限控制、openpyxl Excel 智能导入导出、完整的请假审批工作流（三态流转 + 审批审计链）、学生 CRUD 与自动账号创建。Token 并发刷新队列机制、Promise.all 并行数据加载策略、wx:if 与 hidden 的 DOM 差异选择等工程细节，共同体现了完整的全栈开发思维。感谢大家的聆听，欢迎提问。")

# ========== 保存 ==========
out = r'N:\课程备份\小程序\大作业\小程序演示讲解稿_v4.docx'
doc.save(out)
print(f'[OK] {out}')
