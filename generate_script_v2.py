"""
生成增强版小程序演示讲解稿 — 结合关键代码进行讲解
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ========== 样式工具 ==========
def set_font(run, cn_font, en_font, size_pt, bold=False, color=None):
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.name = en_font
    rPr = run._element.get_or_add_rPr()
    from docx.oxml import OxmlElement
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    if color:
        run.font.color.rgb = color

def add_title(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        if level == 0:
            set_font(run, '黑体', 'Arial', 22, bold=True)
        elif level == 1:
            set_font(run, '黑体', 'Arial', 16, bold=True)
        elif level == 2:
            set_font(run, '黑体', 'Arial', 14, bold=True)
    return h

def add_para(text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74) if indent else Cm(0)
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, '宋体', 'Consolas', 12)
    return p

def add_code(text):
    """行内代码引用 — Consolas 字体 + 蓝色"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, '宋体', 'Consolas', 10.5, color=RGBColor(0x16, 0x77, 0xFF))

def add_bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('● ' + text)
    set_font(run, '宋体', 'Consolas', 12)
    return p

# ================================================================
# 封面
# ================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
run = p.add_run('NETMAXXX 学生管理系统')
set_font(run, '黑体', 'Arial', 26, bold=True, color=RGBColor(0x16, 0x77, 0xFF))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('微信小程序演示讲解稿（技术详解版）')
set_font(run2, '宋体', 'Consolas', 16, color=RGBColor(0x66, 0x66, 0x66))

doc.add_paragraph()

# ================================================================
# 一、项目概述
# ================================================================
add_title('一、项目概述', level=1)

add_para('大家好，我今天展示的项目是 NETMAXXX 学生管理系统。这是一个基于微信小程序平台开发的、面向高校师生日常管理场景的综合信息管理系统。系统支持学生和教师两种角色，涵盖了从学生档案管理、课堂签到、请假审批到数据统计分析、AI 智能对话等完整的校园管理功能模块。')

add_para('系统采用前后端分离架构。前端是微信小程序，使用微信原生框架开发，页面文件由四种类型的文件组成——WXML 模板文件负责界面结构、WXSS 样式文件负责视觉效果、JS 逻辑文件负责交互行为、JSON 配置文件负责页面级设置。所有页面通过 app.json 中的 pages 数组统一注册，由微信客户端按注册顺序加载渲染。')

add_para('后端基于 Django 5.2 框架构建，使用 MySQL 数据库。Django 采用 MVT（Model-View-Template）架构模式，但本项目作为纯 API 后端没有使用 Template 层，而是通过 Django 视图函数直接返回 JsonResponse，本质上是一个 RESTful API 服务。前后端之间的通信全部走 HTTP 协议，JSON 格式传参，JWT Token 做身份认证。整个后端通过 natapp 内网穿透工具将本地 8000 端口映射到公网域名，使小程序能够在任意网络环境下访问后端接口。')

add_para('小程序端的入口文件 app.js 在应用启动时执行 onLaunch 生命周期函数，首先调用 checkLoginStatus 方法检查本地是否存储了有效的 JWT Token。检查逻辑分为三种情况：Token 在有效期内则直接进入主界面并解析用户角色；Token 已过期但 Refresh Token 有效则自动调用刷新接口获取新 Token；没有任何登录凭证则跳转到登录页面。这种启动检查机制保证了用户会话的连续性和安全性。')

# ================================================================
# 二、角色与权限体系
# ================================================================
add_title('二、角色与权限体系', level=1)

add_para('系统设计了学生和教师两种角色，通过 users 表中的 role 字段区分——取值为 student 或 teacher。登录时后端返回的 JWT Token 中嵌入了 role 信息，前端解析 Token 后根据角色展示完全不同的功能界面。这个角色信息同时存储在 app.globalData.userRole 全局变量和本地 Storage 中，确保所有页面都能实时获取。')

add_para('底部导航栏采用了微信小程序的自定义 TabBar 机制实现。与原生 TabBar 不同，自定义 TabBar 是一个普通的自定义组件，定义在 custom-tab-bar 目录下。组件的 JS 逻辑中维护了一个 list 数组，通过过滤 app.globalData.userRole 来动态决定显示哪些标签：教师的标签包括学生信息、数据中心、签到和我的四个入口，学生则显示签到、请假和我的三个入口。每个标签项通过 wx:for 指令遍历 list 数组渲染，使用 image 组件加载图标，通过 data-index 属性标记索引，点击时调用 wx.switchTab 进行页面切换。')

add_para('权限控制的核心在后端。以学生列表接口为例——views.py 中的 get_students 函数会先从请求头中提取 Bearer Token 并解析出用户角色。如果角色是 student，则只返回 user_id 匹配的单条学生记录；如果是 teacher，则返回全量数据。这种在视图函数层面基于角色的数据过滤，从根本上保证了数据安全，前端只做 UI 层面的菜单显隐，真正的权限防线在后端。')

add_code('// 关键代码：Token 解析与角色判断')
add_code('const decoded = api.decodeJWT(token);')
add_code('app.globalData.userRole = decoded.role || \'student\';')
add_code('// 根据角色决定跳转页面')
add_code('const url = decoded.role === \'teacher\' ? \'/pages/student-list/student-list\' : \'/pages/checkin/checkin\';')

# ================================================================
# 三、学生信息管理
# ================================================================
add_title('三、学生信息管理', level=1)

add_para('学生信息管理是整个系统最基础的数据模块。学生档案包含八个字段：学号（SNo，95开头的五位数整数）、姓名（SName，2到5个汉字）、性别（Gender，男/女二选一）、出生日期（Birthday，日期类型）、手机号码（Mobile，11位）、邮箱地址（Email）、家庭住址（Address）和头像照片（Image，存储的是上传后生成的文件名）。在 Django Model 层面，学号被定义为主键——使用 models.IntegerField 并设置 primary_key=True，姓名和性别等字段通过 db_column 参数指定了数据库列名，Gender 字段使用了 choices 参数限制可选值为"男"或"女"。')

add_para('学生列表页的搜索功能通过 Q 查询实现模糊匹配。前端将搜索关键词通过 POST 请求发送到 students/query/ 接口，后端使用 Django 的 Q 对象构建 OR 条件查询——将输入字符串与学号、姓名、性别、手机号、邮箱、地址六个字段逐一进行 icontains 模糊匹配。这样的设计使得教师输入任意一个关键词，无论是姓名、学号还是手机号的一部分，都能快速定位到目标学生。')

add_code('// views.py — 模糊搜索核心代码')
add_code('obj_students = Student.objects.filter(')
add_code('    Q(sno__icontains=data[\'inputstr\']) |')
add_code('    Q(name__icontains=data[\'inputstr\']) |')
add_code('    Q(gender__icontains=data[\'inputstr\']) |')
add_code('    Q(mobile__icontains=data[\'inputstr\']) |')
add_code('    Q(email__icontains=data[\'inputstr\']) |')
add_code('    Q(address__icontains=data[\'inputstr\'])')
add_code(').values()')

add_para('在列表页的 WXML 模板中，每条学生记录用一个 view 卡片展示。卡片左侧使用 image 组件展示头像——当 image 字段有值时，通过 src 属性拼接完整 URL 显示头像图片，使用 mode="aspectFill" 模式保持比例裁剪填充；当没有头像时，显示一个取姓名首字的文字占位符。卡片右侧信息区展示姓名、性别标签、学号和生日，其中性别使用条件类名绑定——如果性别是男则应用 tag-primary 蓝色样式，如果是女则应用 tag-danger 红色样式。卡片右侧的操作按钮仅在教师身份时显示，通过 wx:if="{{userRole === \'teacher\'}}" 控制渲染，包含查看、编辑、删除三个操作入口。')

add_para('Excel 导入导出是教师端的重要功能。导入流程分为三步：首先调用 wx.chooseMessageFile 让教师从微信聊天记录中选择一个 xlsx 或 xls 文件，通过 extension 参数限制文件类型；然后将文件上传到后端 excel/import/ 接口；后端使用 openpyxl 库读取文件内容，通过 read_excel_dict 函数解析数据——该函数支持有表头和无表头两种格式的自动识别，能够将中文列名（如"姓名"）映射到英文字段名（如 name）。导出功能通过 wx.downloadFile 下载后端生成的 Excel 文件，然后利用 wx.shareFileMessage 弹出微信转发面板，教师可以直接将文件发送到聊天；也可以使用 wx.openDocument 在系统查看器中打开。')

add_code('// 前端 — Excel 导入调用')
add_code('wx.chooseMessageFile({')
add_code('  count: 1, type: \'file\', extension: [\'xlsx\', \'xls\'],')
add_code('  success(res) { api.uploadExcel(res.tempFiles[0].path) }')
add_code('});')

# ================================================================
# 四、签到功能（重点）
# ================================================================
add_title('四、签到功能（重点详解）', level=1)

add_para('签到功能是本系统技术实现最丰富的模块，它深度融合了微信小程序的硬件能力——GPS 定位和摄像头扫码。整个模块在单一页面 checkin.js 中实现了教师和学生两种完全不同的交互界面，通过最外层 wx:if="{{userRole === \'teacher\'}}" 和 wx:if="{{userRole === \'student\'}}" 进行分支渲染，共用一套定位工具方法。下面我来详细拆解每个环节的技术实现。')

add_title('4.1 教师发布签到任务', level=2)

add_para('教师创建签到任务的表单包含三个关键交互组件。第一个是签到标题输入框，使用 input 组件，通过 bindinput 事件将输入值实时同步到 taskTitle 数据字段。第二个是签到方式选择器，使用 picker 组件——mode 属性设为 selector 普通选择器，range 绑定 checkinTypes 数组（包含"二维码签到"和"位置签到"两个选项），value 绑定 checkinTypeIndex 索引，用户选择后通过 bindchange 事件回调更新 checkinType 字段。第三个是位置签到专属的范围滑块——使用 slider 组件，min 设为 10 表示最小 10 米，max 设为 2000 表示最大 2000 米，step 设为 10 表示每次拖动步进 10 米，activeColor 设为 #1677FF 品牌蓝色，block-size 设为 20 调整滑块大小。')

add_code('// WXML — 签到类型选择器与范围滑块')
add_code('<picker mode="selector" range="{{checkinTypes}}" value="{{checkinTypeIndex}}" bindchange="onTypeChange">')
add_code('  <view class="picker">{{checkinTypes[checkinTypeIndex]}}</view>')
add_code('</picker>')
add_code('<slider min="10" max="2000" step="10" value="{{rangeValue}}" bindchange="onRangeChange" activeColor="#1677FF" />')

add_para('当签到方式为位置签到时，点击发布按钮前需要先获取教师的位置。系统调用 _getLocation 方法，该方法内部首先通过 wx.getSetting 检查定位权限的授权状态——这是一个关键的设计细节，因为微信的 scope.userLocation 权限有三种状态：true 表示已授权，false 表示用户曾明确拒绝，undefined 表示从未询问过。如果权限为 false，系统会弹出 wx.showModal 对话框引导用户点击"去设置"跳转到系统权限管理页；如果为 undefined，直接调用定位方法，微信会自动弹出授权弹窗。定位调用使用 wx.getLocation，type 参数设为 gcj02 国测局坐标系以适配国内地图服务，isHighAccuracy 设为 true 开启高精度模式，highAccuracyExpireTime 设为 15000 毫秒表示 15 秒后降级为普通精度以节省电量。')

add_code('// JS — 三层权限状态处理逻辑')
add_code('wx.getSetting({')
add_code('  success(res) {')
add_code('    if (res.authSetting[\'scope.userLocation\'] === false) {')
add_code('      wx.showModal({ content: \'请在设置中开启定位权限\', confirmText: \'去设置\' });')
add_code('    } else {')
add_code('      wx.getLocation({ type: \'gcj02\', isHighAccuracy: true });')
add_code('    }')
add_code('  }')
add_code('});')

add_title('4.2 学生扫码签到', level=2)

add_para('二维码签到是两种签到方式中交互最简洁的一种。教师发布任务后，页面上会展示一张二维码图片——通过 image 组件的 src 属性指向后端的二维码生成接口，学生端则看到输入签到码的输入框和"扫描二维码"按钮。学生可以手动输入签到码后点击签到按钮提交，也可以点击扫描按钮调用 wx.scanCode 接口。scanCode 的 onlyFromCamera 参数设为 true，表示仅允许从摄像头扫描，不支持从相册选取——这个细节保证了二维码的物理场景属性，学生必须在课堂上对准教师展示的二维码才能完成签到。扫描成功后在 success 回调中拿到 res.result，将其作为 code 参数连同 task_id 一起通过 POST 请求提交到 checkin/sign/ 接口。')

add_code('// JS — 学生扫码签到')
add_code('wx.scanCode({')
add_code('  onlyFromCamera: true,')
add_code('  success(res) {')
add_code('    api.post(\'checkin/sign/\', { task_id: taskId, type: \'qrcode\', code: res.result });')
add_code('  }')
add_code('});')

add_title('4.3 学生位置签到与地图可视化', level=2)

add_para('位置签到的流程更加复杂，涉及定位获取、地图渲染和用户确认三个步骤。学生点击"获取位置并签到"后，系统先获取学生的经纬度坐标，然后调用 _openStudentMap 方法打开一个全屏地图界面。这个地图界面是整个签到模块最复杂的 UI 组件组合。')

add_para('地图使用微信的 map 组件渲染，这是微信小程序的原生地图组件，背后调用的是腾讯地图服务。组件配置了多个关键属性：latitude 和 longitude 确定地图中心点——我们优先使用教师发布任务时的坐标为地图中心，如果没有则使用学生当前位置；scale 设为 17，对应大约 50 米的比例尺级别；markers 数组包含了两个标记点——红色标记表示"我的位置"，蓝色标记表示"教师位置"。每个标记点使用 callout 属性设置了气泡提示，display 设为 ALWAYS 表示始终显示；label 属性设置了文字标签，通过 anchorX 和 anchorY 控制标签相对于标记点的偏移位置。circles 数组定义了一个半透明圆形覆盖物——以教师坐标为圆心、以签到范围为半径，color 设置边框颜色带透明度，fillColor 设置填充颜色带更高透明度，让学生直观看到签到有效区域。')

add_code('// JS — 构建地图标记点数组')
add_code('markers.push({')
add_code('  id: 0, latitude: studentLat, longitude: studentLng,')
add_code('  width: 32, height: 32,')
add_code('  callout: { content: \'我的位置\', color: \'#E74C3C\', display: \'ALWAYS\' },')
add_code('  label: { content: \'我\', color: \'#E74C3C\', anchorX: 0, anchorY: -12 }')
add_code('});')
add_code('circles.push({')
add_code('  latitude: teacherLat, longitude: teacherLng, radius: range,')
add_code('  color: \'#1677FF66\', fillColor: \'#1677FF1A\', strokeWidth: 2')
add_code('});')

add_para('地图底部叠加了一个半透明的信息面板，展示图例说明和操作按钮。面板内容包括：图例区用三个小色块分别标识"我的位置"（红色圆点）、"教师位置"（蓝色圆点）和"签到范围"（半透明圆环）；确认区域展示"取消"和"确认签到"两个按钮。学生在地图上确认自己的位置在签到范围内后，点击"确认签到"按钮将经纬度数据提交到 checkin/sign/ 接口。')

add_title('4.4 教师查看签到位置与考勤管理', level=2)

add_para('对于位置签到任务，教师可以点击"查看位置"按钮打开一个全屏地图，地图上会同时展示教师参考位置和所有已签到学生的位置。教师端地图使用同样的 map 组件，标记点的 id 按数字编号——id 为 0 的是教师位置，id 从 1 开始的是学生位置。学生标记使用较小的尺寸（24x24 像素）与教师标记（32x32 像素）区分。点击学生标记时触发 bindmarkertap 事件，通过 e.detail.markerId 获取标记 ID，减 1 后从 locationList 数组中取到对应的学生数据，用 wx.showToast 显示学生姓名和学号。')

add_code('// WXML — 教师端地图组件')
add_code('<map id="locationMap" latitude="{{mapLat}}" longitude="{{mapLng}}"')
add_code('  scale="{{mapScale}}" markers="{{mapMarkers}}" circles="{{mapCircles}}"')
add_code('  show-location="{{true}}" bindmarkertap="onMapMarkerTap" />')

add_para('考勤详情面板是整个签到流程的最终环节。教师点击历史签到任务上的"查看考勤"按钮后，系统调用 checkin/attendance/ 接口获取完整考勤数据。后端返回的数据包含三个统计数字——已签到人数 signedCount、未签到人数 unsignedCount 和总人数 total，以及一个详细列表，列表中每个元素包含学生姓名、学号、签到状态和签到时间。未签到学生在列表中以灰色背景标注，教师可以对未签到学生执行"补签"操作——调用 checkin/record/add/ 接口手动添加签到记录，也可以对已签到学生执行"删除"操作移除记录。补签功能的设计解决了学生因手机故障等特殊情况无法正常签到的实际问题。')

# ================================================================
# 五、请假功能（重点）
# ================================================================
add_title('五、请假功能（重点详解）', level=1)

add_para('请假管理是本系统中业务逻辑最完整的工作流模块，实现了从申请提交到审批反馈的完整闭环。数据库层面，Leave 表设计了 12 个字段，不仅记录了请假的基本信息（类型、日期、原因），还通过 student_id、student_name、student_no 三个字段的冗余设计优化了查询性能——教师端可以不需要 JOIN Student 表就读到学生姓名。status 字段使用字符串枚举 pending/approved/rejected 表示审批状态的三态流转。')

add_title('5.1 学生端请假表单', level=2)

add_para('学生端请假页面采用了卡片折叠式布局。顶部是一个可折叠的申请表单卡片，通过 showForm 布尔变量控制表单区域的显示与隐藏——使用 wx:if 指令而非 hidden 属性是为了完全销毁和重建 DOM，在表单收起时清空输入状态。卡片头部使用 bindtap="toggleForm" 切换展开状态，按钮文字通过三元表达式动态显示"收起 ▲"或"展开 ▼"。')

add_para('表单包含四个字段，每个字段有独立的前端校验。请假类型使用 picker 选择器，range 绑定 leaveTypes 数组——包含"事假"、"病假"、"其他"三个选项。开始日期和结束日期使用 mode="date" 的 picker 日期选择器，其中结束日期选择器设置了 start 属性绑定开始日期，微信原生日期选择器会自动将早于开始日期的选项置灰，从 UI 层面杜绝了日期倒挂的输入。请假原因使用 textarea 多行文本组件，maxlength 设为 500 限制最大字符数，当前已输入字符数通过 leaveForm.reason.length 实时显示在右下角。')

add_code('// WXML — 请假表单关键组件')
add_code('<picker mode="selector" range="{{leaveTypes}}" value="{{leaveTypeIndex}}" bindchange="onLeaveTypeChange">')
add_code('<picker mode="date" value="{{leaveForm.startDate}}" bindchange="onStartDateChange">')
add_code('<picker mode="date" value="{{leaveForm.endDate}}" start="{{leaveForm.startDate}}" bindchange="onEndDateChange">')
add_code('<textarea placeholder="请输入请假原因（至少5个字符）" maxlength="500" bindinput="onReasonChange" />')
add_code('<view class="char-count">{{leaveForm.reason.length}}/500</view>')

add_para('前端表单验证通过 validator.validateLeaveForm 函数完成。验证规则包括：请假类型不能为空；开始日期和结束日期不能为空，且结束日期不能早于开始日期；请假原因不能为空且至少 5 个字符。验证不通过时，errors 对象中对应字段会被填充错误信息，WXML 模板通过 wx:if="{{errors.leaveType}}" 条件渲染在字段下方显示红色错误提示文字。所有校验通过后才调用 leave/apply/ 接口提交数据。')

add_para('提交后的请假记录列表支持按审批状态进行本地筛选。状态筛选器使用 picker 组件，选项包括"全部"、"待审批"、"已批准"和"已拒绝"。选择不同的筛选条件后，通过 JavaScript 的 filter 方法在本地对已加载的数据进行过滤——statusIndex 为 0 时 filterStatus 为空字符串，不进行过滤；为 1 到 3 时分别匹配 pending、approved 和 rejected。每条记录以卡片形式展示，顶部用 tag 标签显示请假类型和审批状态，类型和状态的颜色各有不同——病假使用橙色 tag-warning、事假使用灰色 tag-info、待审批使用橙色、已批准使用绿色 tag-success、已拒绝使用红色 tag-danger。')

add_title('5.2 教师端审批工作流', level=2)

add_para('教师端的请假管理页面使用了 wxs 模块来处理模板中的工具函数。wxs（WeiXin Script）是微信小程序特有的脚本语言，运行在视图层而非逻辑层，可以直接在 WXML 中调用。页面引用了一个 wxs 文件——leave-manage.wxs，导出了 isSelected 函数用于判断某条请假记录是否在已选数组中，这个函数在模板中被用于控制复选框的选中样式。使用 wxs 而非在 JS 中预处理数据的好处是减少了数据传递和模板复杂度。')

add_code('// wxs — 模板层工具函数')
add_code('module.exports = {')
add_code('  isSelected: function(arr, id) {')
add_code('    return arr.some(function(item) { return item.id === id; });')
add_code('  }')
add_code('};')

add_para('审批对话框由三部分组成：标题栏根据 approvalType 动态显示"批准请假"或"拒绝请假"；内容区是一个 textarea 输入审批意见（可选填，最多 200 字）；底部是取消和确定两个按钮。确定按钮在提交时携带 leaveId 和 comment 两个参数分别调用 leave/approve/ 或 leave/reject/ 接口。后端在处理审批时，不仅更新 status 字段，还同时写入 teacher_id、approval_comment 和 approval_time 三个字段，完整记录审批操作的执行人和时间，这在后续的责任追溯中非常重要。后端还对请假记录的当前状态进行了二次校验——如果 status 已经不是 pending，则拒绝处理并返回"该请假记录已处理"的提示，防止并发审批导致的数据不一致。')

# ================================================================
# 六、AI 对话功能（重点）
# ================================================================
add_title('六、AI 对话功能（重点详解）', level=1)

add_para('AI 智能对话是本系统的创新亮点，通过接入 DeepSeek 大语言模型为师生提供嵌入业务流程的智能问答服务。整个功能从前端的悬浮面板组件、到工具层的 API 封装、再到后端的模型调用，构成了完整的技术链路。')

add_title('6.1 前端 AI 工具模块', level=2)

add_para('前端在 utils/ai.js 中封装了一个独立的 AI 调用模块。这个模块使用 wx.request 向 DeepSeek 的 API 端点发送 POST 请求，请求体中包含四个参数：model 指定为 deepseek-chat 模型；messages 是对话消息数组，每条消息包含 role（system/user/assistant）和 content 两个字段；temperature 设为 0.7，这是一个在回复创造性和确定性之间的平衡值——数值越高回复越多样，数值越低回复越确定；max_tokens 设为 2000，限制单次回复的最大 Token 数量，既能保证回复完整性又控制了响应时长。模块暴露的 chat 方法支持多轮对话，接收完整的消息历史列表，使得 AI 能够理解对话上下文；ask 方法是对 chat 的简化封装，只传一条 user 消息，适用于单轮问答场景。')

add_code('// utils/ai.js — DeepSeek API 调用')
add_code('function chat(messages) {')
add_code('  return new Promise(function (resolve, reject) {')
add_code('    wx.request({')
add_code('      url: \'https://api.deepseek.com/v1/chat/completions\',')
add_code('      method: \'POST\',')
add_code('      header: { \'Authorization\': \'Bearer \' + _apiKey },')
add_code('      data: { model: \'deepseek-chat\', messages: messages, temperature: 0.7, max_tokens: 2000 },')
add_code('      success: function (res) {')
add_code('        resolve(res.data.choices[0].message.content);')
add_code('      }')
add_code('    });')
add_code('  });')
add_code('}')

add_title('6.2 floating-panel 悬浮面板组件', level=2)

add_para('AI 对话的交互界面实现为一个可复用的自定义组件——floating-panel。组件由四个文件组成：floating-panel.js 是组件的逻辑层，使用 Component 构造器定义，包含 data 数据（对话消息列表、输入框内容、面板展开状态等）和 methods 方法（发送消息、展开/收起面板等）；floating-panel.wxml 是组件的模板层，定义了触发按钮和对话窗口两种 UI 状态；floating-panel.wxss 是组件的样式层，使用组件隔离的样式作用域；floating-panel.json 声明 component: true 表明这是一个自定义组件。')

add_para('面板在默认状态下以一个小圆形浮动按钮的形式定位在页面右下角——使用 position: fixed 固定定位，z-index 设置为较高值确保覆盖在页面内容之上。点击按钮后，面板展开为一个占据屏幕下半部分的对话窗口，使用 animation 动画实现平滑过渡。对话窗口内部是一个经典的聊天界面布局：中间是 scroll-view 可滚动区域展示对话消息，消息气泡采用左右分布——用户消息靠右显示使用蓝色背景，AI 回复靠左显示使用白色背景；底部是固定在窗口下沿的输入区域，包含一个 input 输入框和一个发送按钮。')

add_para('跨页面的 AI 对话持续性通过全局状态管理实现。用户在个人中心点击"AI 对话"菜单后，系统将 app.globalData.aiDialogEnabled 设为 true。每个页面的 onShow 生命周期函数中都会检查这个全局标记——如果为 true 且页面中存在 floating-panel 组件实例，则自动调用组件的 enableTrigger 方法显示悬浮按钮。当用户切换 Tab 页面时，新的页面实例在 onShow 中检测到标记后会自动创建新的 floating-panel 实例，而每个组件实例独立维护自己的对话历史，因此在不同页面间的对话是独立的。')

add_code('// 跨页面 AI 状态同步')
add_code('// profile.js — 用户点击开启 AI 对话')
add_code('enableAiDialog() {')
add_code('  app.globalData.aiDialogEnabled = true;')
add_code('  const fp = this.selectComponent(\'#aiFloatingPanel\');')
add_code('  if (fp) fp.enableTrigger();')
add_code('}')
add_code('')
add_code('// 其他页面 — onShow 时检测并同步')
add_code('onShow() {')
add_code('  if (app.globalData.aiDialogEnabled) {')
add_code('    const fp = this.selectComponent(\'#aiFloatingPanel\');')
add_code('    if (fp) fp.enableTrigger();')
add_code('  }')
add_code('}')

# ================================================================
# 七、数据图表展示（重点）
# ================================================================
add_title('七、数据图表展示（重点详解）', level=1)

add_para('数据可视化看板是教师专属的分析工具，页面加载时通过 Promise.all 并行发起四个 API 请求——getSourceChartData 获取性别分布、getAgeChartData 获取年龄分布、getTypeChartData 获取生日月份分布、getLeaveTrendData 获取近 15 天请假趋势。使用 Promise.all 的优势在于四个请求同时发出，总耗时取决于最慢的那个请求而非四个请求的累加时间。所有请求完成后一次性设置数据并初始化图表，避免数据到达时间不一致导致的渲染闪烁。')

add_code('// JS — 并行加载四个图表数据')
add_code('Promise.all([')
add_code('  api.get(\'getSourceChartData/\'),  // 性别分布')
add_code('  api.get(\'getAgeChartData/\'),     // 年龄分布')
add_code('  api.get(\'getTypeChartData/\'),    // 生日月份')
add_code('  api.get(\'getLeaveTrendData/\')    // 请假趋势')
add_code(']).then(([gender, age, birthday, leave]) => {')
add_code('  // 一次性设置全部数据并渲染图表')
add_code('});')

add_title('7.1 双引擎渲染方案', level=2)

add_para('图表渲染采用了 ECharts 优先、Canvas 2D 降级的双引擎方案。页面加载时首先尝试 require ECharts 库——使用 try/catch 包裹 require 语句，如果 ECharts 已安装则 useECharts 标记为 true，否则进入降级模式。这是一个典型的渐进增强设计：在支持 ECharts 的环境下享受丰富的交互特效（tooltip 提示框、legend 图例切换、动画过渡），在不支持的环境下仍然能用原生 Canvas 2D 呈现核心数据。')

add_para('ECharts 的集成通过 ec-canvas 组件完成。这是一个社区适配的桥接组件，其核心机制是在组件的 ready 生命周期中通过 wx.createSelectorQuery 获取 Canvas 节点，然后调用 ECharts 的 init 方法将 Canvas 实例初始化为 ECharts 图表对象。每个图表的配置通过 ec 属性传入组件——ec 对象包含一个 onInit 回调函数，该函数接收 canvas、width、height 和 dpr（设备像素比）四个参数，在函数内部调用 echarts.init 创建图表实例，通过 setOption 设置图表配置项，最后返回图表实例供组件管理生命周期。以性别饼图为例，其 ECharts 配置中 tooltip 的 trigger 设为 item 表示点击扇区时触发提示，legend 图例置于底部居中，series 系列类型为 pie，radius 设为数组 ["45%", "70%"] 生成环形饼图，center 调整圆心位置为图例预留空间，label 的 formatter 设为 "{b}\n{d}%" 在每个扇区显示名称和百分比。')

add_code('// JS — ECharts 图表初始化配置（性别饼图）')
add_code('chart.setOption({')
add_code('  tooltip: { trigger: \'item\' },')
add_code('  legend: { bottom: 0, left: \'center\' },')
add_code('  series: [{')
add_code('    type: \'pie\',')
add_code('    radius: [\'45%\', \'70%\'],  // 环形饼图')
add_code('    center: [\'50%\', \'45%\'],')
add_code('    data: genderData,')
add_code('    label: { formatter: \'{b}\\n{d}%\' }  // 名称+百分比')
add_code('  }]')
add_code('});')

add_title('7.2 Canvas 2D 降级方案', level=2)

add_para('降级方案使用微信小程序原生 Canvas 2D API 手绘三种图表类型——饼图、柱状图和折线图。这是一个展示 Canvas 编程能力的典型案例。以饼图为例，drawPieChart 函数首先通过 wx.createSelectorQuery 获取 Canvas 节点——与旧版 Canvas 不同，新版 Canvas 2D 使用 fields({ node: true, size: true }) 获取节点和尺寸信息，然后通过 canvas.getContext("2d") 获取 2D 绘图上下文。绘图前需要适配设备像素比——canvas.width 和 canvas.height 设为逻辑尺寸乘以 dpr，然后 ctx.scale(dpr, dpr) 缩放画布，保证在 Retina 屏幕上文字和图形仍然清晰锐利。')

add_para('饼图的核心绘制逻辑：先计算数据总和，然后遍历每个数据项，用 item.value 除以总和得到该扇区对应的弧度角 sliceAngle。以圆心 (cx, cy) 为起点，使用 ctx.arc 方法从 startAngle 绘制到 startAngle + sliceAngle，每次绘制前用 ctx.beginPath 开始新路径，绘制后用 ctx.fill 填充颜色。完成一个扇区后，计算该扇区中间角度的方向，在该方向上距圆心 labelDist 处绘制数据标签文字。标签内容包括类别名称、数值和百分比。柱状图和折线图的绘制遵循同样的 Canvas 2D 编程模式——先绘制网格参考线，再根据数据计算坐标绘制图形主体，最后标注数值。')

add_code('// JS — Canvas 2D 降级方案：饼图扇区绘制')
add_code('let startAngle = -Math.PI / 2;  // 从12点钟方向开始')
add_code('data.forEach((item, i) => {')
add_code('  const sliceAngle = (item.value / total) * 2 * Math.PI;')
add_code('  ctx.beginPath();')
add_code('  ctx.moveTo(cx, cy);')
add_code('  ctx.arc(cx, cy, radius, startAngle, startAngle + sliceAngle);')
add_code('  ctx.closePath();')
add_code('  ctx.fillStyle = colors[i % colors.length];')
add_code('  ctx.fill();')
add_code('  startAngle += sliceAngle;')
add_code('});')

add_para('WXML 模板中，ECharts 和 Canvas 降级方案通过 wx:if 和 wx:else 条件渲染实现互斥：当 useECharts 为 true 时渲染 ec-canvas 组件，组件上绑定 id、ec 配置对象等属性；当 useECharts 为 false 时渲染原生 canvas 标签，type 必须设为 "2d" 指定使用新版 Canvas 2D API，通过 id 属性在 JS 中获取节点引用。两种方案的图表都放在 chart-card 卡片容器中，卡片带有圆角、阴影和标题，在 charts-grid 网格布局中自适应排列——小图表占一半宽度，全宽图表占满整行。')

add_title('7.3 后端数据统计', level=2)

add_para('图表数据的后端统计逻辑各具特色。性别统计最简单——使用 Django ORM 的 filter 和 count 方法分别统计 gender 为"男"和"女"的记录数，返回 [{name: "男", value: 人数}, {name: "女", value: 人数}] 格式。年龄统计需要先查出所有学生，遍历每个学生的 birthday 字段计算实际年龄——用当前年份减出生年份，再根据月日比较修正未到生日的情况。生日月份统计使用 Django ORM 的 annotate 方法添加计算字段——ExtractMonth 函数从 birthday 中提取月份数字，然后按月份分组计数，最后映射为一月到十二月的中文名称，并确保每月都有数据（即使人数为 0）。请假趋势统计查询 Leave 表中 apply_time 在最近 15 天内的记录，使用 TruncDate 按天截断时间，按天分组 count 计数，最后对没有记录的日期补 0 以保证折线图的连续性。')

# ================================================================
# 八、Token 管理与自动刷新
# ================================================================
add_title('八、Token 管理与自动刷新机制', level=1)

add_para('系统的认证安全通过 JWT 双 Token 机制实现，这是前后端分离项目中常用的认证方案。Access Token 有效期设为 30 分钟——这个时间是在安全性和用户体验之间的平衡，太短用户频繁掉线，太长则泄露后风险窗口过大。Refresh Token 有效期 7 天，仅用于获取新的 Access Token，不参与业务请求。')

add_para('前端 api.js 中的 request 函数封装了一套自动刷新机制。当请求返回 HTTP 401 状态码时——这通常意味着 Access Token 已过期——函数并非立刻跳转登录页，而是先检查是否已有刷新操作正在进行中。这里使用了一个巧妙的设计：isRefreshing 布尔标记和 refreshQueue 等待队列。如果当前没有正在进行的刷新操作，则置 isRefreshing 为 true 并调用 doRefreshToken 函数；如果已有刷新操作正在进行中，则将重试当前请求的回调函数推入 refreshQueue 队列。当 doRefreshToken 成功获取新 Token 后，遍历 refreshQueue 执行所有排队的回调，每个回调以携带新 Token 的方式重新发起原始请求。这个机制解决了多个请求同时在 401 后各自触发刷新导致 Token 混乱的并发问题。如果刷新也失败了——意味着 Refresh Token 也过期了——则清除本地所有登录凭证并跳转到登录页面。')

add_code('// api.js — 401 拦截与并发刷新控制')
add_code('if (res.statusCode === 401 && !options._retry) {')
add_code('  if (isRefreshing) {')
add_code('    refreshQueue.push(() => {')
add_code('      request({ ...options, _retry: true }).then(resolve).catch(reject);')
add_code('    });')
add_code('    return;')
add_code('  }')
add_code('  isRefreshing = true;')
add_code('  doRefreshToken().then(() => {')
add_code('    refreshQueue.forEach(cb => cb());')
add_code('    refreshQueue = [];')
add_code('    doRequest();  // 重试原始请求')
add_code('  });')
add_code('}')

# ================================================================
# 九、全局样式与组件化
# ================================================================
add_title('九、全局样式体系与组件化设计', level=1)

add_para('系统的样式架构采用了 CSS 变量 + 工具类 + 组件样式三层体系。全局样式文件 app.wxss 中定义了一套完整的设计变量——使用 CSS 自定义属性（也叫 CSS 变量）在 page 选择器上声明，包括品牌色 --color-primary（#1677FF 蓝色）、成功色 --color-success（#00B42A 绿色）、警告色 --color-warning（#FF7D00 橙色）、危险色 --color-danger（#F53F3F 红色）等语义化颜色变量，以及 --radius-sm/md/lg 圆角变量、--spacing-xs/sm/md/lg 间距变量、--font-xs/sm/md/lg/xl 字号变量。所有页面和组件通过 var() 函数引用这些变量，当需要调整主题色时只需修改一处。')

add_para('全局工具类提供了 flex 布局、文本对齐、颜色、间距等高频使用的 CSS 片段。例如 .flex-between 组合了 display: flex、align-items: center 和 justify-content: space-between 三个属性，在 WXML 中一行 class="flex-between" 就能实现水平两端对齐的弹性布局。.tag 系列样式定义了标签组件的基础样式和五种颜色变体——.tag-primary 使用浅蓝底蓝字、.tag-success 浅绿底绿字、.tag-warning 浅橙底橙字、.tag-danger 浅红底红字、.tag-info 浅灰底灰字。这套标签体系在整个系统中统一使用，无论是请假类型标签、签到状态标签、性别标签还是审批状态标签，都通过复用相同的 class 名称保持一致的外观。')

add_para('在组件化方面，系统抽取了三个可复用的自定义组件。floating-panel 是 AI 对话悬浮面板，封装了完整的对话 UI 和交互逻辑。ec-canvas 是 ECharts 的 Canvas 桥接组件，解决了 ECharts 在微信小程序中的初始化问题。empty-state 是空状态占位组件，在数据列表为空时显示统一的空状态提示，接收 title 和 description 两个属性。自定义组件通过 JSON 文件中的 "component": true 声明，页面通过 usingComponents 配置引用，组件与页面之间通过 properties 向下传递数据和通过 triggerEvent 向上发射事件进行通信。')

# ================================================================
# 十、总结
# ================================================================
add_title('十、总结', level=1)

add_para('NETMAXXX 学生管理系统是一个功能完整、架构清晰、细节考究的微信小程序项目。从前端来看，它综合运用了 WXML 的数据绑定与条件渲染、WXSS 的 CSS 变量与工具类体系、WXS 的视图层脚本、自定义组件的封装与复用、微信小程序地图和扫码等硬件 API 的调用、ECharts 和 Canvas 2D 的双引擎图表渲染。从后端来看，它实现了 Django ORM 的数据建模与查询、JWT 双 Token 的认证与自动刷新、openpyxl 的 Excel 读写、bcrypt 的密码加密存储、基于角色的接口级权限控制。')

add_para('系统的四个核心功能模块各有其技术深度：签到模块深度融合了 GPS 定位的三层权限处理、地图组件的 markers/circles 可视化配置和扫码 API 的摄像头调用；请假模块实现了完整的三态审批工作流，前后端均有数据校验和状态流转控制；AI 对话通过悬浮面板组件和 DeepSeek API 集成，实现了跨页面的智能助手；数据图表则通过 ECharts 优先加 Canvas 降级的双引擎方案，在效果和兼容性之间取得平衡。')

add_para('在工程细节上，Token 的并发刷新队列机制、Promise.all 的并行数据加载策略、WXS 的模板层工具函数、CSS 变量的统一主题管理、以及权限的三层防护——前端路由、后端接口、数据库关联——都体现了完整的工程思维。感谢大家的聆听，如果有什么问题，欢迎提问。')

# ========== 保存 ==========
output_path = r'N:\课程备份\小程序\大作业\小程序演示讲解稿_v2.docx'
doc.save(output_path)
print(f'[OK] Saved to: {output_path}')
