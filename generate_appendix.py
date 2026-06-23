"""
生成附录 Word 文档 — NETMAXXX 小程序前端 + Django 后端核心代码
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ========== 样式设置 ==========
style = doc.styles['Normal']
font = style.font
font.name = '等线'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

# 代码块样式（灰色背景小字）
def add_code_block(doc, code_text, language=""):
    """添加代码块，用灰色底纹 + Courier New 字体"""
    for line in code_text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(13)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        # 灰色底纹
        from docx.oxml import OxmlElement
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'F5F5F5')
        shd.set(qn('w:val'), 'clear')
        p.paragraph_format.element.get_or_add_pPr().append(shd)

def add_file_header(doc, filepath, description=""):
    """添加文件名标题"""
    p = doc.add_paragraph()
    run = p.add_run(f"📄 {filepath}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x16, 0x77, 0xFF)
    if description:
        run2 = p.add_run(f"  — {description}")
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x86, 0x90, 0x9C)

def read_file(path):
    """读取文件，处理编码"""
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    except:
        return f"[无法读取文件: {path}]"

# ========== 标题页 ==========
title = doc.add_heading('附录：系统核心代码', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('NETMAXXX 学生管理系统 — 微信小程序 + Django 后端')
doc.add_paragraph('')

# ================================================================
# 第一部分：微信小程序前端
# ================================================================
doc.add_heading('一、微信小程序前端', level=1)

PROJ = r'M:\WXML\EMAXXX\miniprogram'

# ---------- 1. 入口文件 ----------
doc.add_heading('1. 入口文件', level=2)

add_file_header(doc, 'app.js', '全局入口逻辑 — 登录状态检查、Token刷新、全局数据')
add_code_block(doc, read_file(os.path.join(PROJ, 'app.js')))

add_file_header(doc, 'app.json', '全局配置 — 页面路由、TabBar、权限声明')
add_code_block(doc, read_file(os.path.join(PROJ, 'app.json')))

add_file_header(doc, 'app.wxss', '全局样式 — 设计变量、通用组件、工具类')
add_code_block(doc, read_file(os.path.join(PROJ, 'app.wxss')))

# ---------- 2. 工具模块 ----------
doc.add_heading('2. 工具模块 (utils/)', level=2)

add_file_header(doc, 'utils/api.js', 'API 请求封装 — Token管理、自动刷新、请求拦截、文件上传')
add_code_block(doc, read_file(os.path.join(PROJ, 'utils', 'api.js')))

add_file_header(doc, 'utils/validator.js', '表单验证 — 学号/姓名/手机/邮箱/日期校验规则')
add_code_block(doc, read_file(os.path.join(PROJ, 'utils', 'validator.js')))

add_file_header(doc, 'utils/ai.js', 'AI 对话工具 — DeepSeek API 封装')
add_code_block(doc, read_file(os.path.join(PROJ, 'utils', 'ai.js')))

# ---------- 3. 自定义 TabBar ----------
doc.add_heading('3. 自定义 TabBar', level=2)

add_file_header(doc, 'custom-tab-bar/index.js', '自定义底部导航 — 角色感知的TabBar')
add_code_block(doc, read_file(os.path.join(PROJ, 'custom-tab-bar', 'index.js')))

add_file_header(doc, 'custom-tab-bar/index.wxml', 'TabBar 模板')
add_code_block(doc, read_file(os.path.join(PROJ, 'custom-tab-bar', 'index.wxml')))

add_file_header(doc, 'custom-tab-bar/index.json', 'TabBar 配置')
add_code_block(doc, read_file(os.path.join(PROJ, 'custom-tab-bar', 'index.json')))

# ---------- 4. 登录页 ----------
doc.add_heading('4. 登录/注册页 (pages/login/)', level=2)

add_file_header(doc, 'pages/login/login.wxml', '登录注册模板 — 学生/教师双角色登录注册表单')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'login', 'login.wxml')))

add_file_header(doc, 'pages/login/login.wxss', '登录页样式 — 渐变背景、模式切换、角色选择器')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'login', 'login.wxss')))

add_file_header(doc, 'pages/login/login.js', '登录注册逻辑 — 表单验证、JWT解析、角色路由')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'login', 'login.js')))

add_file_header(doc, 'pages/login/login.json', '登录页配置')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'login', 'login.json')))

# ---------- 5. 学生列表页 ----------
doc.add_heading('5. 学生列表页 (pages/student-list/)', level=2)

add_file_header(doc, 'pages/student-list/student-list.wxml', '学生列表模板 — 搜索、分页、批量操作、Excel导入导出')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'student-list', 'student-list.wxml')))

add_file_header(doc, 'pages/student-list/student-list.wxss', '学生列表样式')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'student-list', 'student-list.wxss')))

add_file_header(doc, 'pages/student-list/student-list.js', '学生列表逻辑 — CRUD、搜索、分页、批量删除、Excel操作')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'student-list', 'student-list.js')))

# ---------- 6. 学生表单页 ----------
doc.add_heading('6. 学生表单页 (pages/student-form/)', level=2)

add_file_header(doc, 'pages/student-form/student-form.wxml', '学生表单模板 — 添加/编辑/查看三模式、头像上传、字段校验')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'student-form', 'student-form.wxml')))

add_file_header(doc, 'pages/student-form/student-form.wxss', '学生表单样式')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'student-form', 'student-form.wxss')))

add_file_header(doc, 'pages/student-form/student-form.js', '学生表单逻辑 — 三模式切换、学号异步校验、头像上传')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'student-form', 'student-form.js')))

# ---------- 7. 签到模块 ----------
doc.add_heading('7. 签到模块 (pages/checkin/)', level=2)

add_file_header(doc, 'pages/checkin/checkin.wxml', '签到模板 — 教师发布/管理签到 + 学生扫码/定位签到 + 地图 + 考勤')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'checkin', 'checkin.wxml')))

add_file_header(doc, 'pages/checkin/checkin.wxss', '签到样式 — 地图全屏、位置面板、考勤弹窗')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'checkin', 'checkin.wxss')))

add_file_header(doc, 'pages/checkin/checkin.js', '签到逻辑 — 教师发布/结束/查看位置 + 学生签到/扫码/定位 + 考勤管理')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'checkin', 'checkin.js')))

# ---------- 8. 请假申请（学生端）----------
doc.add_heading('8. 请假申请页 - 学生端 (pages/leave-apply/)', level=2)

add_file_header(doc, 'pages/leave-apply/leave-apply.wxml', '请假申请模板 — 类型选择、日期范围、原因填写、记录查看')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'leave-apply', 'leave-apply.wxml')))

add_file_header(doc, 'pages/leave-apply/leave-apply.wxss', '请假申请样式')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'leave-apply', 'leave-apply.wxss')))

add_file_header(doc, 'pages/leave-apply/leave-apply.js', '请假申请逻辑 — 表单提交、状态筛选、分页')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'leave-apply', 'leave-apply.js')))

# ---------- 9. 请假管理（教师端）----------
doc.add_heading('9. 请假管理页 - 教师端 (pages/leave-manage/)', level=2)

add_file_header(doc, 'pages/leave-manage/leave-manage.wxml', '请假管理模板 — 筛选、审批、批量删除')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'leave-manage', 'leave-manage.wxml')))

add_file_header(doc, 'pages/leave-manage/leave-manage.wxss', '请假管理样式')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'leave-manage', 'leave-manage.wxss')))

add_file_header(doc, 'pages/leave-manage/leave-manage.js', '请假管理逻辑 — 搜索筛选、批准/拒绝、批量删除')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'leave-manage', 'leave-manage.js')))

# ---------- 10. 数据看板 ----------
doc.add_heading('10. 数据看板 (pages/charts/)', level=2)

add_file_header(doc, 'pages/charts/charts.wxml', '图表模板 — 教师看4类统计图 / 学生请假模块')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'charts', 'charts.wxml')))

add_file_header(doc, 'pages/charts/charts.wxss', '图表样式')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'charts', 'charts.wxss')))

add_file_header(doc, 'pages/charts/charts.js', '图表逻辑 — ECharts初始化 + Canvas降级方案(饼/柱/折线)')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'charts', 'charts.js')))

# ---------- 11. 个人中心 ----------
doc.add_heading('11. 个人中心 (pages/profile/)', level=2)

add_file_header(doc, 'pages/profile/profile.wxml', '个人中心模板 — 角色菜单枢纽')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'profile', 'profile.wxml')))

add_file_header(doc, 'pages/profile/profile.wxss', '个人中心样式')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'profile', 'profile.wxss')))

add_file_header(doc, 'pages/profile/profile.js', '个人中心逻辑 — 头像加载、菜单导航、AI对话、退出登录')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'profile', 'profile.js')))

# ---------- 12. 教师信息 ----------
doc.add_heading('12. 讲师信息页 (pages/teacher-info/)', level=2)

add_file_header(doc, 'pages/teacher-info/teacher-info.wxml', '教师信息模板 — 查看/编辑切换、头像上传')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'teacher-info', 'teacher-info.wxml')))

add_file_header(doc, 'pages/teacher-info/teacher-info.wxss', '教师信息样式')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'teacher-info', 'teacher-info.wxss')))

add_file_header(doc, 'pages/teacher-info/teacher-info.js', '教师信息逻辑 — 数据加载、编辑保存、头像上传')
add_code_block(doc, read_file(os.path.join(PROJ, 'pages', 'teacher-info', 'teacher-info.js')))

# ---------- 13. 组件 ----------
doc.add_heading('13. 自定义组件 (components/)', level=2)

add_file_header(doc, 'components/floating-panel/floating-panel.js', 'AI 悬浮面板 — DeepSeek 对话界面')
add_code_block(doc, read_file(os.path.join(PROJ, 'components', 'floating-panel', 'floating-panel.js')))

add_file_header(doc, 'components/floating-panel/floating-panel.wxml', 'AI 悬浮面板模板')
add_code_block(doc, read_file(os.path.join(PROJ, 'components', 'floating-panel', 'floating-panel.wxml')))

add_file_header(doc, 'components/ec-canvas/ec-canvas.js', 'ECharts Canvas 组件 — 微信小程序适配')
add_code_block(doc, read_file(os.path.join(PROJ, 'components', 'ec-canvas', 'ec-canvas.js')))

add_file_header(doc, 'components/empty-state/empty-state.js', '空状态组件')
add_code_block(doc, read_file(os.path.join(PROJ, 'components', 'empty-state', 'empty-state.js')))

add_file_header(doc, 'components/empty-state/empty-state.wxml', '空状态模板')
add_code_block(doc, read_file(os.path.join(PROJ, 'components', 'empty-state', 'empty-state.wxml')))

# ================================================================
# 第二部分：Django 后端核心代码
# ================================================================
doc.add_heading('二、Django 后端核心代码', level=1)

BE = r'N:\课程备份\小程序\大作业\Backend-source'

doc.add_heading('1. 数据模型 (apps/student/models.py)', level=2)
add_file_header(doc, 'apps/student/models.py', 'Django ORM 模型 — Student、Teacher、User、Leave')
add_code_block(doc, read_file(os.path.join(BE, 'apps', 'student', 'models.py')))

doc.add_heading('2. 认证视图 (apps/student/auth_views.py)', level=2)
add_file_header(doc, 'apps/student/auth_views.py', 'JWT 认证 — 登录/注册/Token刷新')
add_code_block(doc, read_file(os.path.join(BE, 'apps', 'student', 'auth_views.py')))

doc.add_heading('3. 学生管理视图 (apps/student/views.py)', level=2)
add_file_header(doc, 'apps/student/views.py', '学生 CRUD — 搜索/增删改/Excel导入导出/图表统计')
add_code_block(doc, read_file(os.path.join(BE, 'apps', 'student', 'views.py')))

doc.add_heading('4. 请假管理视图 (apps/student/leave_views.py)', level=2)
add_file_header(doc, 'apps/student/leave_views.py', '请假管理 — 申请/审批/删除')
add_code_block(doc, read_file(os.path.join(BE, 'apps', 'student', 'leave_views.py')))

doc.add_heading('5. 教师管理视图 (apps/student/teacher_views.py)', level=2)
add_file_header(doc, 'apps/student/teacher_views.py', '教师信息 — 获取/更新')
add_code_block(doc, read_file(os.path.join(BE, 'apps', 'student', 'teacher_views.py')))

doc.add_heading('6. URL 路由 (StudentV4BE/urls.py)', level=2)
add_file_header(doc, 'StudentV4BE/urls.py', 'API 路由配置')
add_code_block(doc, read_file(os.path.join(BE, 'StudentV4BE', 'urls.py')))

doc.add_heading('7. Django 设置 (StudentV4BE/settings.py)', level=2)
add_file_header(doc, 'StudentV4BE/settings.py', '数据库/CORS/中间件配置')
add_code_block(doc, read_file(os.path.join(BE, 'StudentV4BE', 'settings.py')))

# ================================================================
# 保存
# ================================================================
output_path = r'N:\课程备份\小程序\大作业\附录_系统核心代码.docx'
doc.save(output_path)
print(f'[OK] Document saved to: {output_path}')
