"""
修改附录文档：
1. 删除所有 📄 图标
2. 文件名设为三级标题：Consolas 小四(12pt)
3. 保存为新文件
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'N:\课程备份\小程序\大作业\附录_系统核心代码_格式化.docx'
DST = r'N:\课程备份\小程序\大作业\附录_系统核心代码_最终版.docx'

doc = Document(SRC)

H3_SIZE = Pt(12)  # 小四

def set_h3_font(run):
    """设置为三级标题：Consolas 小四 加粗"""
    run.font.size = H3_SIZE
    run.bold = True
    run.font.name = 'Consolas'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')

def is_blue_bold_header(para):
    """判断是否为蓝色加粗文件名头"""
    for run in para.runs:
        if run.bold:
            try:
                if run.font.color and run.font.color.rgb:
                    r, g, b = run.font.color.rgb
                    if r < 80 and b > 200:
                        return True
            except:
                pass
    return False

def remove_emoji(text):
    """删除 📄 等 emoji 图标"""
    result = []
    for ch in text:
        # 过滤掉 emoji 和特殊符号（📄 U+1F4C4 等）
        cp = ord(ch)
        if cp >= 0x1F000 and cp <= 0x1FFFF:
            continue
        if cp >= 0x2700 and cp <= 0x27BF:
            continue
        if cp >= 0x2600 and cp <= 0x26FF:
            continue
        result.append(ch)
    return ''.join(result).strip()

modified = 0

for para in doc.paragraphs:
    if not para.text.strip():
        continue

    if is_blue_bold_header(para):
        # 清除首行缩进
        para.paragraph_format.first_line_indent = Cm(0)

        for run in para.runs:
            # 删除 emoji
            run.text = remove_emoji(run.text)
            # 取消蓝色，改为默认黑色
            run.font.color.rgb = None
            # 统一字体（描述文字也改为三级标题格式）
            set_h3_font(run)

        # 处理：如果描述部分（灰色文字）存在，也统一格式
        # 让整个段落所有 run 都用三级标题格式
        all_runs_h3 = True
        for run in para.runs:
            set_h3_font(run)

        modified += 1

# 也检查普通段落中的 emoji
for para in doc.paragraphs:
    for run in para.runs:
        new_text = remove_emoji(run.text)
        if new_text != run.text:
            run.text = new_text

doc.save(DST)
print(f'[OK] Saved: {DST}')
print(f'  {modified} file headers updated: emoji removed, Consolas 小四(12pt)')
