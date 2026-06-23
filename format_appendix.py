"""
格式化附录 Word 文档
- 一级标题：黑体 二号(22pt)
- 二级标题：黑体 三号(16pt)
- 三级标题（文件名）：宋体 小三(15pt)
- 正文：中文宋体/英文Consolas 小四(12pt)，首行缩进2字符
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import copy

doc_path = r'N:\课程备份\小程序\大作业\附录_系统核心代码.docx'
doc = Document(doc_path)

# ========== 字号对照 ==========
# 二号 = 22pt, 三号 = 16pt, 小三 = 15pt, 小四 = 12pt

def set_run_font(run, font_name_cn, font_name_en, size_pt, bold=False, color=None):
    """统一设置 run 的字体属性"""
    run.font.size = Pt(size_pt)
    run.bold = bold
    # 设置中文字体
    run.font.name = font_name_en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name_cn)
    rFonts.set(qn('w:ascii'), font_name_en)
    rFonts.set(qn('w:hAnsi'), font_name_en)
    if color:
        run.font.color.rgb = color

def apply_heading_style(paragraph, font_cn, font_en, size_pt):
    """设置段落为标题格式"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    for run in paragraph.runs:
        set_run_font(run, font_cn, font_en, size_pt, bold=True)

def apply_body_style(paragraph):
    """设置正文格式"""
    paragraph.paragraph_format.first_line_indent = Cm(0.74)  # 约2个字符
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(16)
    for run in paragraph.runs:
        # 保持原有的颜色等属性，只改字体和字号
        old_color = run.font.color.rgb if run.font.color and run.font.color.rgb else None
        set_run_font(run, '宋体', 'Consolas', 12, bold=False, color=old_color)

def is_code_block(paragraph):
    """判断是否为代码块（灰色底纹 + Courier New 小字）"""
    if not paragraph.runs:
        return False
    run = paragraph.runs[0]
    if run.font.name == 'Courier New' and run.font.size and run.font.size < Pt(11):
        return True
    return False

def is_file_header(paragraph):
    """判断是否为文件名标题（蓝色加粗文本）"""
    if not paragraph.runs:
        return False
    run = paragraph.runs[0]
    if run.bold and run.font.size and run.font.size <= Pt(12):
        try:
            if run.font.color.rgb:
                r, g, b = run.font.color.rgb
                if r < 50 and g > 100 and b > 200:  # 蓝色系
                    return True
        except:
            pass
    return False

# ========== 遍历所有段落 ==========
heading1_count = 0
heading2_count = 0
heading3_count = 0
body_count = 0
code_count = 0

for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue

    # 检测段落样式
    style_name = para.style.name if para.style else ''

    # 一级标题：Heading 1
    if 'Heading 1' in style_name or 'heading 1' in style_name:
        apply_heading_style(para, '黑体', 'Arial', 22)
        heading1_count += 1
        continue

    # 二级标题：Heading 2
    if 'Heading 2' in style_name or 'heading 2' in style_name:
        apply_heading_style(para, '黑体', 'Arial', 16)
        heading2_count += 1
        continue

    # 三级标题：Heading 3 或 文件头部（蓝色加粗）
    if 'Heading 3' in style_name or 'heading 3' in style_name:
        apply_heading_style(para, '宋体', 'Consolas', 15)
        heading3_count += 1
        continue

    if is_file_header(para):
        apply_heading_style(para, '宋体', 'Consolas', 15)
        heading3_count += 1
        continue

    # 跳过代码块（保持原有 Courier New 小字格式）
    if is_code_block(para):
        code_count += 1
        continue

    # 其余视为正文
    apply_body_style(para)
    body_count += 1

# ========== 也处理表格中的文本 ==========
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                apply_body_style(para)

# ========== 保存 ==========
doc.save(doc_path)
print(f'[OK] Formatting applied:')
print(f'  Heading 1 (黑体 22pt): {heading1_count}')
print(f'  Heading 2 (黑体 16pt): {heading2_count}')
print(f'  Heading 3 (宋体 15pt): {heading3_count}')
print(f'  Body (宋体/Consolas 12pt): {body_count}')
print(f'  Code blocks (unchanged): {code_count}')
print(f'  Saved to: {doc_path}')
